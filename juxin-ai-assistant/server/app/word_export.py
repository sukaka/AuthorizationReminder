from __future__ import annotations

from datetime import date
from io import BytesIO
import re
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from app.document_templates.company_style import (
    COMPANY_REQUIRED_SECTIONS,
    COMPANY_WORD_STYLE,
    FACT_RISK_CONTROL_ITEMS,
    FINAL_REVIEW_SECTIONS,
)

CHINESE_NUMERALS = (
    "",
    "一",
    "二",
    "三",
    "四",
    "五",
    "六",
    "七",
    "八",
    "九",
    "十",
)


def render_generation_docx(
    *,
    title: str,
    task_name: str,
    department: str,
    author: str,
    output: str,
    version: str,
) -> bytes:
    document = Document()
    _configure_document(document, title=title, version=version)
    heading_numbers = [0, 0, 0]

    _add_cover_page(
        document,
        title=title,
        task_name=task_name,
        department=department,
        author=author,
        version=version,
    )
    _add_revision_table(
        document,
        author=author,
        version=version,
        heading_numbers=heading_numbers,
    )

    _add_heading(document, f"《{title}》", 0)
    _render_markdown_blocks(document, output, heading_numbers)
    _add_missing_company_required_sections(document, output, heading_numbers)
    _add_missing_fact_risk_control_items(document, output)
    _add_missing_final_review_sections(document, output, heading_numbers)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _configure_document(document: Document, *, title: str, version: str) -> None:
    page_style = COMPANY_WORD_STYLE["page"]
    font_style = COMPANY_WORD_STYLE["font"]
    brand_style = COMPANY_WORD_STYLE["brand"]
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(page_style["top_margin_cm"])
    section.bottom_margin = Cm(page_style["bottom_margin_cm"])
    section.left_margin = Cm(page_style["left_margin_cm"])
    section.right_margin = Cm(page_style["right_margin_cm"])
    section.header_distance = Cm(page_style["header_distance_cm"])
    section.footer_distance = Cm(page_style["footer_distance_cm"])

    styles = document.styles
    _set_style_font(styles["Normal"], font_style["body"], 11)
    _set_style_font(styles["Title"], font_style["heading"], 22, bold=True)
    for style_name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        _set_style_font(styles[style_name], font_style["heading"], size, bold=True)

    header_paragraph = section.header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_text_run(header_paragraph, f"{brand_style['name']}｜{title}｜{version}", bold=True)
    _set_paragraph_bottom_border(header_paragraph, color=brand_style["header_line_color"])

    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_text_run(
        footer_paragraph,
        f"{brand_style['company']}｜{brand_style['confidentiality']} / 客户项目交付文档｜第 ",
    )
    _add_field(footer_paragraph, "PAGE")
    _add_text_run(footer_paragraph, " 页 / 共 ")
    _add_field(footer_paragraph, "NUMPAGES")
    _add_text_run(footer_paragraph, " 页")


def _add_cover_page(
    document: Document,
    *,
    title: str,
    task_name: str,
    department: str,
    author: str,
    version: str,
) -> None:
    font_style = COMPANY_WORD_STYLE["font"]
    brand_style = COMPANY_WORD_STYLE["brand"]
    document.add_paragraph()
    document.add_paragraph()

    brand = document.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_text_run(brand, brand_style["name"], font=font_style["heading"], size=16, bold=True)

    document.add_paragraph()

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_text_run(title_paragraph, f"《{title}》", font=font_style["heading"], size=24, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_text_run(subtitle, f"{brand_style['name']}公司级 AI 生成文档", font=font_style["body"], size=14)

    document.add_paragraph()
    metadata = [
        ("文档名称", title),
        ("客户或部门", department or "待确认"),
        ("项目或系统", task_name or "待确认"),
        ("版本", version or "待确认"),
        ("编制单位", department or "待确认"),
        ("编制人员", author or "待确认"),
        ("日期", date.today().isoformat()),
        ("保密级别", brand_style["confidentiality"]),
    ]
    table = document.add_table(rows=len(metadata), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row, (label, value) in zip(table.rows, metadata):
        _write_cell(row.cells[0], label, bold=True, shading=brand_style["table_header_fill"])
        _write_cell(row.cells[1], value)

    document.add_page_break()


def _add_revision_table(
    document: Document,
    *,
    author: str,
    version: str,
    heading_numbers: list[int],
) -> None:
    brand_style = COMPANY_WORD_STYLE["brand"]
    _add_heading(document, "修订记录", 1, heading_numbers)
    table = document.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ("版本", "日期", "修订内容", "修订人", "备注")
    values = (
        version or "待确认",
        date.today().isoformat(),
        "初始生成",
        author or "待确认",
        "待确认",
    )
    for cell, text in zip(table.rows[0].cells, headers):
        _write_cell(cell, text, bold=True, shading=brand_style["table_header_fill"])
    for cell, text in zip(table.rows[1].cells, values):
        _write_cell(cell, text)

    document.add_paragraph()


def _render_markdown_blocks(
    document: Document,
    markdown: str,
    heading_numbers: list[int],
) -> None:
    lines = markdown.splitlines()
    index = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        _add_body_paragraph(document, " ".join(item.strip() for item in paragraph_buffer).strip())
        paragraph_buffer.clear()

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            if index < len(lines):
                index += 1
            _add_code_block(document, code_lines)
            continue

        if _is_pipe_table_line(stripped):
            flush_paragraph()
            table_lines: list[str] = []
            while index < len(lines) and _is_pipe_table_line(lines[index].strip()):
                table_lines.append(lines[index].strip())
                index += 1
            _add_pipe_table(document, table_lines)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            _add_heading(
                document,
                heading.group(2).strip(),
                len(heading.group(1)),
                heading_numbers,
            )
            index += 1
            continue

        list_item = re.match(r"^(\d+[.)]|[-*+])\s+(.+)$", stripped)
        if list_item:
            flush_paragraph()
            style = "List Number" if list_item.group(1)[0].isdigit() else "List Bullet"
            paragraph = document.add_paragraph(style=style)
            _add_text_run(paragraph, list_item.group(2).strip())
            index += 1
            continue

        paragraph_buffer.append(line)
        index += 1

    flush_paragraph()


def _add_missing_final_review_sections(
    document: Document,
    output: str,
    heading_numbers: list[int],
) -> None:
    for heading in FINAL_REVIEW_SECTIONS:
        if heading in output:
            continue
        _add_heading(document, heading, 1, heading_numbers)
        _add_body_paragraph(document, "待确认")


def _add_missing_company_required_sections(
    document: Document,
    output: str,
    heading_numbers: list[int],
) -> None:
    for heading in COMPANY_REQUIRED_SECTIONS:
        if heading in output:
            continue
        _add_heading(document, heading, 1, heading_numbers)
        if heading == "风险与注意事项":
            for label, value in FACT_RISK_CONTROL_ITEMS:
                _add_body_paragraph(document, f"{label}：{value}")
            _add_body_paragraph(document, "待确认事项：待确认。")
        elif heading in {"表格清单或结果统计", "交付物或附件"}:
            _add_body_paragraph(document, "暂无")
        else:
            _add_body_paragraph(document, "待确认")


def _add_missing_fact_risk_control_items(document: Document, output: str) -> None:
    existing_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for label, value in FACT_RISK_CONTROL_ITEMS:
        if label not in output and label not in existing_text:
            _add_body_paragraph(document, f"{label}：{value}")


def _add_heading(
    document: Document,
    text: str,
    level: int,
    heading_numbers: list[int] | None = None,
) -> Paragraph:
    font_style = COMPANY_WORD_STYLE["font"]
    if level == 0:
        paragraph = document.add_paragraph(style="Title")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph = document.add_paragraph(style=f"Heading {min(level, 3)}")
        if heading_numbers is not None:
            text = _number_heading(text, min(level, 3), heading_numbers)
    _add_text_run(paragraph, text, font=font_style["heading"], bold=True)
    return paragraph


def _number_heading(text: str, level: int, heading_numbers: list[int]) -> str:
    text = _strip_heading_number(text, level)
    _advance_heading_number(level, heading_numbers)
    if level == 1:
        return f"{_to_chinese_number(heading_numbers[0])}、{text}"
    if level == 2:
        return f"{heading_numbers[1]}. {text}"
    return f"（{heading_numbers[2]}）{text}"


def _advance_heading_number(level: int, heading_numbers: list[int]) -> None:
    if level == 1:
        heading_numbers[0] += 1
        heading_numbers[1] = 0
        heading_numbers[2] = 0
    elif level == 2:
        heading_numbers[1] += 1
        heading_numbers[2] = 0
    else:
        heading_numbers[2] += 1


def _strip_heading_number(text: str, level: int) -> str:
    if level == 1:
        return re.sub(r"^[一二三四五六七八九十百千万零〇两]+、\s*", "", text)
    if level == 2:
        return re.sub(r"^\d+[.)]\s*", "", text)
    return re.sub(r"^（\d+）\s*", "", text)


def _to_chinese_number(value: int) -> str:
    if 0 < value < len(CHINESE_NUMERALS):
        return CHINESE_NUMERALS[value]
    if 10 < value < 20:
        return f"十{CHINESE_NUMERALS[value - 10]}"
    if value == 20:
        return "二十"
    if 20 < value < 100:
        tens, ones = divmod(value, 10)
        return f"{CHINESE_NUMERALS[tens]}十{CHINESE_NUMERALS[ones]}"
    return str(value)


def _add_body_paragraph(document: Document, text: str) -> Paragraph:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    _add_text_run(paragraph, text)
    return paragraph


def _add_code_block(document: Document, lines: Iterable[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.5)
        _add_text_run(paragraph, line, font="Courier New")


def _add_pipe_table(document: Document, lines: list[str]) -> None:
    rows = [_split_pipe_row(line) for line in lines]
    rows = [row for row in rows if row and not _is_separator_row(row)]
    if not rows:
        for line in lines:
            _add_body_paragraph(document, line)
        return

    max_columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=max_columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row_values in enumerate(rows):
        for col_index in range(max_columns):
            value = row_values[col_index] if col_index < len(row_values) else ""
            value = value.strip() or "待确认"
            _write_cell(
                table.cell(row_index, col_index),
                value,
                bold=row_index == 0,
                shading=COMPANY_WORD_STYLE["brand"]["table_header_fill"] if row_index == 0 else None,
            )

    document.add_paragraph()


def _is_pipe_table_line(line: str) -> bool:
    return line.startswith("|") and line.endswith("|") and line.count("|") >= 2


def _split_pipe_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_separator_row(row: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in row)


def _write_cell(cell: _Cell, text: str, *, bold: bool = False, shading: str | None = None) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if shading:
        _shade_cell(cell, shading)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.clear()
    _add_text_run(paragraph, text, bold=bold)


def _add_text_run(
    paragraph: Paragraph,
    text: str,
    *,
    font: str | None = None,
    size: int | None = None,
    bold: bool = False,
) -> Run:
    run = paragraph.add_run(text)
    font = font or COMPANY_WORD_STYLE["font"]["body"]
    run.font.name = font
    run.font.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    _set_run_east_asia_font(run, font)
    return run


def _set_style_font(style, font_name: str, size: int, *, bold: bool = False) -> None:
    style.font.name = font_name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_run_east_asia_font(run: Run, font_name: str) -> None:
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_paragraph_bottom_border(paragraph: Paragraph, *, color: str) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    border = paragraph_properties.find(qn("w:pBdr"))
    if border is None:
        border = OxmlElement("w:pBdr")
        paragraph_properties.append(border)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    border.append(bottom)


def _shade_cell(cell: _Cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _add_field(paragraph: Paragraph, instruction: str) -> None:
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instruction} "
    run._r.append(instr_text)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    run._r.append(placeholder)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)

import hashlib
import uuid as uuid_lib
from io import BytesIO

from docx import Document
from fastapi import HTTPException, UploadFile
from sqlalchemy import select
from sqlalchemy.orm import Session

from .crypto import ContentCipher, EncryptedPayload
from .knowledge_files import _extract_blocks, _validated_document_archive
from .models import GenerationAttachment, Task


MAX_ATTACHMENT_MB = 100
MAX_ATTACHMENT_BYTES = MAX_ATTACHMENT_MB * 1024 * 1024
SUPPORTED_TEXT_SUFFIXES = {".txt", ".md", ".docx", ".xlsx", ".pptx", ".pdf"}
UNSUPPORTED_TYPE_MESSAGE = "当前仅支持 docx、xlsx、pptx、txt、md"
PDF_UNSUPPORTED_MESSAGE = "PDF 文本提取将在下一步启用；扫描件暂不支持 OCR"


def _safe_file_name(raw_name: str | None) -> str:
    file_name = (raw_name or "").replace("\\", "/").rsplit("/", 1)[-1].strip()
    if not file_name:
        raise HTTPException(status_code=422, detail="文件名不能为空")
    if len(file_name) > 255:
        raise HTTPException(status_code=422, detail="文件名不能超过 255 个字符")
    return file_name


def _file_suffix(file_name: str) -> str:
    dot_index = file_name.rfind(".")
    return file_name[dot_index:].lower() if dot_index >= 0 else ""


def _parse_docx_text(data: bytes) -> str:
    try:
        # DOCX is a ZIP container too; validate entry count, expansion ratio,
        # and total uncompressed size before python-docx opens it.
        with _validated_document_archive(data, file_kind="DOCX"):
            pass
        document = Document(BytesIO(data))
        parts: list[str] = []
        parts.extend(
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        )
        for table in document.tables:
            for row in table.rows:
                parts.append(
                    " | ".join(
                        cell.text.strip() or "待确认"
                        for cell in row.cells
                    )
                )
        return "\n".join(parts).strip()
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=422, detail="DOCX 文件无法解析") from exc


def _extract_text(file_name: str, data: bytes) -> str:
    suffix = _file_suffix(file_name)
    if suffix not in SUPPORTED_TEXT_SUFFIXES:
        raise HTTPException(status_code=415, detail=UNSUPPORTED_TYPE_MESSAGE)
    if suffix == ".pdf":
        raise HTTPException(status_code=422, detail=PDF_UNSUPPORTED_MESSAGE)
    if suffix in {".txt", ".md"}:
        try:
            return data.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise HTTPException(
                status_code=422,
                detail="文本附件必须使用 UTF-8 编码",
            ) from exc
    if suffix in {".xlsx", ".pptx"}:
        return "\n\n".join(block.text for block in _extract_blocks(file_name, data) if block.text.strip())

    return _parse_docx_text(data)


async def create_attachment(
    db: Session,
    sso_user_id: str,
    task_uuid: str,
    file: UploadFile,
    cipher: ContentCipher,
    key_version: str,
) -> tuple[GenerationAttachment, int]:
    task = db.scalar(
        select(Task).where(Task.uuid == task_uuid, Task.status == "ACTIVE")
    )
    if task is None:
        raise HTTPException(status_code=404, detail="任务不存在或未启用")

    file_name = _safe_file_name(file.filename)
    content = await file.read(MAX_ATTACHMENT_BYTES + 1)
    if len(content) > MAX_ATTACHMENT_BYTES:
        raise HTTPException(status_code=413, detail=f"附件大小不能超过 {MAX_ATTACHMENT_MB} MB")

    extracted_text = _extract_text(file_name, content)

    attachment_uuid = str(uuid_lib.uuid4())
    encrypted = cipher.encrypt_json(
        {"text": extracted_text},
        attachment_uuid.encode(),
    )
    attachment = GenerationAttachment(
        uuid=attachment_uuid,
        sso_user_id=sso_user_id,
        task_id=task.id,
        file_name=file_name,
        file_type=(file.content_type or "application/octet-stream")[:128],
        file_size=len(content),
        content_sha256=hashlib.sha256(content).hexdigest(),
        extracted_text_ciphertext=encrypted.ciphertext,
        extracted_text_nonce=encrypted.nonce,
        key_version=key_version,
        status="READY",
        error_code="",
    )
    db.add(attachment)
    db.flush()
    db.refresh(attachment)
    return attachment, len(extracted_text)


def load_owned_attachment_texts(
    db: Session,
    sso_user_id: str,
    task_id: int,
    attachment_uuids: list[str],
    cipher: ContentCipher,
) -> list[tuple[GenerationAttachment, str]]:
    if not attachment_uuids:
        return []

    records = list(db.scalars(
        select(GenerationAttachment).where(
            GenerationAttachment.uuid.in_(attachment_uuids),
            GenerationAttachment.sso_user_id == sso_user_id,
            GenerationAttachment.task_id == task_id,
            GenerationAttachment.status == "READY",
        )
    ))
    records_by_uuid = {record.uuid: record for record in records}
    if len(records_by_uuid) != len(attachment_uuids):
        raise HTTPException(status_code=404, detail="附件不存在或无权访问")

    loaded: list[tuple[GenerationAttachment, str]] = []
    for attachment_uuid in attachment_uuids:
        record = records_by_uuid[attachment_uuid]
        payload = cipher.decrypt_json(
            EncryptedPayload(
                record.extracted_text_ciphertext,
                record.extracted_text_nonce,
            ),
            record.uuid.encode(),
        )
        loaded.append((record, str(payload.get("text", ""))))
    return loaded

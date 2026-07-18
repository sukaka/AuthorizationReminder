import hashlib
import re
import uuid as uuid_lib
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from fastapi import HTTPException
from pypdf import PdfReader
from pypdf.errors import PdfReadError
from sqlalchemy import delete, select
from sqlalchemy.orm import Session

from .crypto import ContentCipher, EncryptedPayload
from .knowledge_search import EmbeddingService
from .models import KnowledgeChunk, KnowledgeFile


MAX_KNOWLEDGE_FILE_MB = 100
MAX_KNOWLEDGE_FILE_BYTES = MAX_KNOWLEDGE_FILE_MB * 1024 * 1024
MAX_ARCHIVE_ENTRIES = 2_000
MAX_ARCHIVE_ENTRY_UNCOMPRESSED_BYTES = 32 * 1024 * 1024
MAX_ARCHIVE_UNCOMPRESSED_BYTES = 128 * 1024 * 1024
MAX_ARCHIVE_COMPRESSION_RATIO = 100
MAX_PDF_PAGES = 500
MAX_EXTRACTED_TEXT_CHARS = 5 * 1024 * 1024
MAX_EXTRACTED_BLOCKS = 5_000
MAX_KNOWLEDGE_METADATA_TEXT_LENGTH = 255
SUPPORTED_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp"}
SUPPORTED_KNOWLEDGE_SUFFIXES = {".txt", ".md", ".docx", ".xlsx", ".pptx", ".pdf", *SUPPORTED_IMAGE_SUFFIXES}
UNSUPPORTED_KNOWLEDGE_TYPE_MESSAGE = "当前版本暂不支持该文件类型，请上传 pdf、docx、xlsx、pptx、txt、md、png、jpg、jpeg 或 webp 文件。"
HEADING_PATTERN = re.compile(r"^(#{1,6}\s+|[一二三四五六七八九十]+、|\d+[.．、]|\（\d+\）)\S+")
SPECIAL_TERM_PATTERN = re.compile(
    r"(?i)\b(?:CVE-\d{4}-\d{4,}|[A-Z]{2,}[-_]?[A-Z0-9]{2,}[-_]?\d*|GB/T\s*\d+|ISO\s*\d+|[0-9]{1,3}(?:\.[0-9]{1,3}){3}|[0-9]{2,5})\b"
)


@dataclass(frozen=True)
class ChunkDraft:
    text: str
    chunk_index: int
    section_title: str
    page_number: int | None = None
    section_path: str = ""
    page_or_sheet: str = ""
    chunk_type: str = "text"
    metadata: dict | None = None


@dataclass(frozen=True)
class ParsedBlock:
    text: str
    section_path: str = ""
    page_or_sheet: str = ""
    chunk_type: str = "paragraph"
    metadata: dict | None = None


def _safe_file_name(raw_name: str | None) -> str:
    file_name = (raw_name or "").replace("\\", "/").rsplit("/", 1)[-1].strip()
    if not file_name:
        raise HTTPException(status_code=422, detail="文件名不能为空")
    if len(file_name) > 255:
        raise HTTPException(status_code=422, detail="文件名不能超过 255 个字符")
    return file_name


def _limit_metadata_text(value: str | None) -> str:
    return (value or "").strip()[:MAX_KNOWLEDGE_METADATA_TEXT_LENGTH]


def _file_suffix(file_name: str) -> str:
    dot_index = file_name.rfind(".")
    return file_name[dot_index:].lower() if dot_index >= 0 else ""


def _is_relative_to(path: Path, parent: Path) -> bool:
    try:
        path.relative_to(parent)
    except ValueError:
        return False
    return True


def _storage_subdir(usage_type: str) -> Path:
    if usage_type == "official_knowledge":
        return Path("knowledge") / "original"
    if usage_type == "session_attachment":
        return Path("user_uploads") / "session_attachments"
    return Path("user_uploads") / "personal_references"


def _persist_original_file(
    *,
    storage_root: str | None,
    usage_type: str,
    file_uuid: str,
    file_name: str,
    content: bytes,
) -> tuple[str, str]:
    if not storage_root:
        return "", ""
    root = Path(storage_root).expanduser().resolve()
    suffix = _file_suffix(file_name)
    stored_file_name = f"{file_uuid}{suffix}"
    target_dir = (root / _storage_subdir(usage_type)).resolve()
    target_path = (target_dir / stored_file_name).resolve()
    if not _is_relative_to(target_path, root):
        raise HTTPException(status_code=500, detail="知识文件存储路径无效")
    target_dir.mkdir(parents=True, exist_ok=True)
    target_path.write_bytes(content)
    return stored_file_name, str(target_path)


def _decode_utf8_text(data: bytes, *, message: str) -> str:
    try:
        return data.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise HTTPException(status_code=422, detail=message) from exc


def _validated_document_archive(data: bytes, *, file_kind: str) -> ZipFile:
    """Open an Office archive only after bounding its decompression cost."""
    archive = ZipFile(BytesIO(data))
    entries = archive.infolist()
    if len(entries) > MAX_ARCHIVE_ENTRIES:
        archive.close()
        raise HTTPException(status_code=422, detail=f"{file_kind} 文件包含过多压缩条目")
    total_size = 0
    for entry in entries:
        if entry.flag_bits & 0x1:
            archive.close()
            raise HTTPException(status_code=422, detail=f"{file_kind} 文件不支持加密压缩条目")
        total_size += entry.file_size
        if entry.file_size > MAX_ARCHIVE_ENTRY_UNCOMPRESSED_BYTES:
            archive.close()
            raise HTTPException(status_code=422, detail=f"{file_kind} 文件解压后的单个条目过大")
        if entry.file_size > 1024 * 1024 and entry.compress_size > 0 and entry.file_size / entry.compress_size > MAX_ARCHIVE_COMPRESSION_RATIO:
            archive.close()
            raise HTTPException(status_code=422, detail=f"{file_kind} 文件压缩比异常")
    if total_size > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
        archive.close()
        raise HTTPException(status_code=422, detail=f"{file_kind} 文件解压后的内容过大")
    return archive


def _validate_extracted_blocks(blocks: list[ParsedBlock]) -> list[ParsedBlock]:
    if len(blocks) > MAX_EXTRACTED_BLOCKS:
        raise HTTPException(status_code=422, detail="文件可提取内容分段过多")
    if sum(len(block.text) for block in blocks) > MAX_EXTRACTED_TEXT_CHARS:
        raise HTTPException(status_code=422, detail="文件可提取文本过大")
    return blocks


def _clean_row(row: list[str]) -> list[str]:
    cleaned_cells = [cell.strip() for cell in row]
    while cleaned_cells and not cleaned_cells[-1]:
        cleaned_cells.pop()
    return cleaned_cells


def _markdown_table(rows: list[list[str]]) -> str:
    cleaned_rows = [_clean_row(row) for row in rows]
    cleaned_rows = [row for row in cleaned_rows if row]
    if not cleaned_rows:
        return ""
    column_count = max(len(row) for row in cleaned_rows)
    normalized = [
        row + ["待确认"] * (column_count - len(row))
        for row in cleaned_rows
    ]
    header = normalized[0]
    separator = ["---"] * column_count
    body = normalized[1:]
    table_rows = [header, separator, *body]
    return "\n".join("| " + " | ".join(cell or "待确认" for cell in row) + " |" for row in table_rows)


def _row_record_text(headers: list[str], row: list[str]) -> str:
    if not headers:
        return " | ".join(cell or "待确认" for cell in row)
    padded = row + ["待确认"] * max(0, len(headers) - len(row))
    pairs = [
        f"{header or f'字段{index + 1}'}={padded[index] or '待确认'}"
        for index, header in enumerate(headers)
    ]
    return "；".join(pairs)


def _xml_texts(root: ET.Element) -> list[str]:
    return [
        (node.text or "").strip()
        for node in root.findall(".//{*}t")
        if (node.text or "").strip()
    ]


def _keywords(text: str) -> list[str]:
    keywords = []
    for match in SPECIAL_TERM_PATTERN.findall(text):
        normalized = re.sub(r"\s+", " ", match.strip())
        if normalized and normalized not in keywords:
            keywords.append(normalized)
    return keywords[:50]


def _parse_docx_blocks(data: bytes) -> list[ParsedBlock]:
    try:
        with _validated_document_archive(data, file_kind="DOCX"):
            pass
        document = Document(BytesIO(data))
    except HTTPException:
        raise
    except BadZipFile as exc:
        raise HTTPException(status_code=422, detail="DOCX 文件无法解析") from exc
    except Exception as exc:
        raise HTTPException(status_code=422, detail="DOCX 文件无法解析") from exc

    blocks: list[ParsedBlock] = []
    section_stack: list[str] = []
    current_section = ""
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        heading_level = 0
        if style_name.startswith("heading"):
            match = re.search(r"(\d+)", style_name)
            heading_level = int(match.group(1)) if match else 1
        elif HEADING_PATTERN.match(text):
            heading_level = 1
        if heading_level:
            section_stack = section_stack[: max(heading_level - 1, 0)]
            section_stack.append(text.lstrip("#").strip())
            current_section = " / ".join(section_stack)
            blocks.append(ParsedBlock(text=text, section_path=current_section, chunk_type="heading"))
            continue
        blocks.append(ParsedBlock(text=text, section_path=current_section, chunk_type="paragraph"))

    for table_index, table in enumerate(document.tables, start=1):
        rows = [
            [cell.text.strip() for cell in row.cells]
            for row in table.rows
        ]
        table_text = _markdown_table(rows)
        if table_text:
            section = current_section or f"表格 {table_index}"
            blocks.append(ParsedBlock(text=table_text, section_path=section, chunk_type="table"))
    return blocks


def _xlsx_shared_strings(archive: ZipFile) -> list[str]:
    try:
        shared_strings = archive.read("xl/sharedStrings.xml")
    except KeyError:
        return []
    try:
        root = ET.fromstring(shared_strings)
    except ET.ParseError as exc:
        raise HTTPException(status_code=422, detail="XLSX 共享字符串无法解析") from exc
    values: list[str] = []
    for item in root.findall(".//{*}si"):
        text_parts = [
            text_node.text or ""
            for text_node in item.findall(".//{*}t")
        ]
        values.append("".join(text_parts))
    return values


def _xlsx_cell_text(cell: ET.Element, shared_strings: list[str]) -> str:
    cell_type = cell.attrib.get("t", "")
    if cell_type == "inlineStr":
        text_node = cell.find(".//{*}t")
        return (text_node.text or "") if text_node is not None else ""

    value_node = cell.find("{*}v")
    if value_node is None or value_node.text is None:
        return ""
    raw_value = value_node.text.strip()
    if cell_type == "s":
        try:
            return shared_strings[int(raw_value)]
        except (ValueError, IndexError) as exc:
            raise HTTPException(status_code=422, detail="XLSX 单元格共享字符串索引无效") from exc
    return raw_value


def _xlsx_sheet_names(archive: ZipFile) -> dict[str, str]:
    try:
        workbook_root = ET.fromstring(archive.read("xl/workbook.xml"))
        rels_root = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
    except (KeyError, ET.ParseError):
        return {}
    relationships = {
        rel.attrib.get("Id", ""): rel.attrib.get("Target", "")
        for rel in rels_root.findall(".//{*}Relationship")
    }
    names: dict[str, str] = {}
    for sheet in workbook_root.findall(".//{*}sheet"):
        rel_id = sheet.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id", "")
        target = relationships.get(rel_id, "")
        if not target:
            continue
        if not target.startswith("xl/"):
            target = "xl/" + target.lstrip("/")
        names[target] = sheet.attrib.get("name", "") or target.rsplit("/", 1)[-1].removesuffix(".xml")
    return names


def _parse_xlsx_blocks(data: bytes) -> list[ParsedBlock]:
    try:
        with _validated_document_archive(data, file_kind="XLSX") as archive:
            shared_strings = _xlsx_shared_strings(archive)
            sheet_names = _xlsx_sheet_names(archive)
            worksheet_names = sorted(
                name
                for name in archive.namelist()
                if name.startswith("xl/worksheets/") and name.endswith(".xml")
            )
            blocks: list[ParsedBlock] = []
            for worksheet_name in worksheet_names:
                try:
                    root = ET.fromstring(archive.read(worksheet_name))
                except ET.ParseError as exc:
                    raise HTTPException(status_code=422, detail="XLSX 工作表无法解析") from exc
                sheet_name = sheet_names.get(
                    worksheet_name,
                    worksheet_name.rsplit("/", 1)[-1].removesuffix(".xml"),
                )
                rows: list[list[str]] = []
                for row in root.findall(".//{*}row"):
                    cleaned = _clean_row([
                        _xlsx_cell_text(cell, shared_strings)
                        for cell in row.findall("{*}c")
                    ])
                    if cleaned:
                        rows.append(cleaned)
                if not rows:
                    continue
                headers = rows[0]
                lines = [f"Sheet：{sheet_name}", f"表头：{' | '.join(headers)}", _markdown_table(rows)]
                for row in rows[1:]:
                    lines.append("记录：" + _row_record_text(headers, row))
                blocks.append(
                    ParsedBlock(
                        text="\n".join(line for line in lines if line).strip(),
                        section_path=sheet_name,
                        page_or_sheet=sheet_name,
                        chunk_type="sheet_rows",
                        metadata={
                            "sheet_name": sheet_name,
                            "headers": headers,
                        },
                    )
                )
    except BadZipFile as exc:
        raise HTTPException(status_code=422, detail="XLSX 文件无法解析") from exc
    except KeyError as exc:
        raise HTTPException(status_code=422, detail="XLSX 文件结构不完整") from exc
    return blocks


def _slide_index(path: str) -> int:
    match = re.search(r"slide(\d+)\.xml$", path, re.IGNORECASE)
    return int(match.group(1)) if match else 0


def _parse_pptx_blocks(data: bytes) -> list[ParsedBlock]:
    try:
        with _validated_document_archive(data, file_kind="PPTX") as archive:
            slide_names = sorted(
                (
                    name
                    for name in archive.namelist()
                    if name.startswith("ppt/slides/slide") and name.endswith(".xml")
                ),
                key=_slide_index,
            )
            note_text_by_index: dict[int, str] = {}
            for note_name in archive.namelist():
                if not note_name.startswith("ppt/notesSlides/notesSlide") or not note_name.endswith(".xml"):
                    continue
                try:
                    note_root = ET.fromstring(archive.read(note_name))
                except ET.ParseError:
                    continue
                note_text_by_index[_slide_index(note_name)] = "\n".join(_xml_texts(note_root))

            blocks: list[ParsedBlock] = []
            for slide_name in slide_names:
                slide_number = _slide_index(slide_name)
                try:
                    root = ET.fromstring(archive.read(slide_name))
                except ET.ParseError as exc:
                    raise HTTPException(status_code=422, detail="PPTX 幻灯片无法解析") from exc
                texts = _xml_texts(root)
                if not texts:
                    continue
                title = texts[0]
                body = "\n".join(texts[1:]).strip()
                note_text = note_text_by_index.get(slide_number, "").strip()
                parts = [
                    f"幻灯片 {slide_number}：{title}",
                    body,
                    f"备注：{note_text}" if note_text else "",
                ]
                blocks.append(
                    ParsedBlock(
                        text="\n".join(part for part in parts if part).strip(),
                        section_path=title,
                        page_or_sheet=f"幻灯片 {slide_number}",
                        chunk_type="slide",
                        metadata={"slide_index": slide_number},
                    )
                )
    except BadZipFile as exc:
        raise HTTPException(status_code=422, detail="PPTX 文件无法解析") from exc
    except KeyError as exc:
        raise HTTPException(status_code=422, detail="PPTX 文件结构不完整") from exc
    return blocks


def _parse_pdf_blocks(data: bytes) -> list[ParsedBlock]:
    try:
        reader = PdfReader(BytesIO(data))
    except (PdfReadError, ValueError, OSError) as exc:
        raise HTTPException(status_code=400, detail="PDF 文件解析失败，请确认文件未损坏。") from exc

    blocks: list[ParsedBlock] = []
    if len(reader.pages) > MAX_PDF_PAGES:
        raise HTTPException(status_code=422, detail="PDF 页数超过处理上限")
    for page_index, page in enumerate(reader.pages, start=1):
        try:
            text = (page.extract_text() or "").strip()
        except (PdfReadError, ValueError, OSError):
            text = ""
        if not text:
            continue
        page_title = f"第 {page_index} 页"
        blocks.append(
            ParsedBlock(
                text=text,
                section_path=page_title,
                page_or_sheet=page_title,
                chunk_type="pdf_page",
                metadata={"page_number": page_index, "page_index": page_index},
            )
        )

    if not blocks:
        raise HTTPException(status_code=400, detail="PDF 未提取到可用文本，请上传可复制文本的 PDF。")
    return blocks


def _parse_text_blocks(file_name: str, data: bytes) -> list[ParsedBlock]:
    text = _decode_utf8_text(data, message="文本知识文件必须使用 UTF-8 编码")
    is_markdown = _file_suffix(file_name) == ".md"
    blocks: list[ParsedBlock] = []
    current_section = ""
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_lines
        if current_lines:
            blocks.append(
                ParsedBlock(
                    text="\n".join(current_lines).strip(),
                    section_path=current_section,
                    chunk_type="markdown" if is_markdown else "paragraph",
                )
            )
            current_lines = []

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            flush()
            continue
        if HEADING_PATTERN.match(line):
            flush()
            current_section = line.lstrip("#").strip()
            current_lines = [line]
            continue
        current_lines.append(line)
    flush()
    return blocks


def _extract_blocks(file_name: str, data: bytes) -> list[ParsedBlock]:
    suffix = _file_suffix(file_name)
    if suffix not in SUPPORTED_KNOWLEDGE_SUFFIXES:
        raise HTTPException(status_code=415, detail=UNSUPPORTED_KNOWLEDGE_TYPE_MESSAGE)
    if suffix in SUPPORTED_IMAGE_SUFFIXES:
        valid_signature = (
            (suffix == ".png" and data.startswith(b"\x89PNG\r\n\x1a\n"))
            or (suffix in {".jpg", ".jpeg"} and data.startswith(b"\xff\xd8\xff"))
            or (suffix == ".webp" and len(data) >= 12 and data[:4] == b"RIFF" and data[8:12] == b"WEBP")
        )
        if not valid_signature:
            raise HTTPException(status_code=422, detail="图片内容与文件扩展名不匹配")
        display_name = file_name.rsplit(".", 1)[0].replace("_", " ").replace("-", " ").strip()
        blocks = [ParsedBlock(
            text=f"图片资料：{display_name or file_name}\n文件名：{file_name}",
            section_path="图片资料",
            chunk_type="image",
            metadata={"media_type": suffix.lstrip("."), "image_asset": True},
        )]
    elif suffix in {".txt", ".md"}:
        blocks = _parse_text_blocks(file_name, data)
    elif suffix == ".xlsx":
        blocks = _parse_xlsx_blocks(data)
    elif suffix == ".pptx":
        blocks = _parse_pptx_blocks(data)
    elif suffix == ".pdf":
        blocks = _parse_pdf_blocks(data)
    else:
        blocks = _parse_docx_blocks(data)
    return _validate_extracted_blocks(blocks)


def _section_blocks(text: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    current_title = ""
    current_lines: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if HEADING_PATTERN.match(line):
            if current_lines:
                blocks.append((current_title, "\n".join(current_lines)))
            current_title = line
            current_lines = [line]
            continue
        current_lines.append(line)
    if current_lines:
        blocks.append((current_title, "\n".join(current_lines)))
    if blocks:
        return blocks
    stripped = text.strip()
    return [("", stripped)] if stripped else []


def chunk_text(
    text: str,
    *,
    target_chars: int = 1000,
    max_chars: int = 1500,
    overlap_chars: int = 150,
) -> list[ChunkDraft]:
    if max_chars <= 0:
        raise ValueError("max_chars must be greater than 0")
    safe_overlap = max(0, min(overlap_chars, max_chars - 1))
    safe_target = max(1, min(target_chars, max_chars))
    chunks: list[ChunkDraft] = []
    for section_title, content in _section_blocks(text):
        start = 0
        while start < len(content):
            end = min(start + max_chars, len(content))
            chunk_body = content[start:end].strip()
            if chunk_body:
                chunks.append(
                    ChunkDraft(
                        text=chunk_body,
                        chunk_index=len(chunks),
                        section_title=section_title,
                        section_path=section_title,
                        chunk_type="text",
                        metadata={"keywords": _keywords(chunk_body)},
                    )
                )
            if end >= len(content):
                break
            next_start = end - safe_overlap
            if end - start > safe_target and safe_overlap == 0:
                next_start = end
            start = max(start + 1, next_start)
    return chunks


def chunk_blocks(
    blocks: list[ParsedBlock],
    *,
    target_chars: int = 1000,
    max_chars: int = 1500,
    overlap_chars: int = 150,
) -> list[ChunkDraft]:
    chunks: list[ChunkDraft] = []
    for block in blocks:
        block_text = block.text.strip()
        if not block_text:
            continue
        block_metadata = dict(block.metadata or {})
        block_page_number = block_metadata.get("page_number")
        if not isinstance(block_page_number, int):
            block_page_number = None
        if block.chunk_type in {"table", "sheet_rows", "slide"} and len(block_text) <= max_chars:
            metadata = dict(block_metadata)
            metadata["keywords"] = _keywords(block_text)
            chunks.append(
                ChunkDraft(
                    text=block_text,
                    chunk_index=len(chunks),
                    section_title=block.section_path,
                    page_number=block_page_number,
                    section_path=block.section_path,
                    page_or_sheet=block.page_or_sheet,
                    chunk_type=block.chunk_type,
                    metadata=metadata,
                )
            )
            continue
        for draft in chunk_text(
            block_text,
            target_chars=target_chars,
            max_chars=max_chars,
            overlap_chars=overlap_chars,
        ):
            section_path = block.section_path or draft.section_path or draft.section_title
            metadata = dict(block_metadata)
            metadata["keywords"] = _keywords(draft.text)
            chunks.append(
                ChunkDraft(
                    text=draft.text,
                    chunk_index=len(chunks),
                    section_title=section_path,
                    page_number=draft.page_number or block_page_number,
                    section_path=section_path,
                    page_or_sheet=block.page_or_sheet,
                    chunk_type=block.chunk_type,
                    metadata=metadata,
                )
            )
    return chunks


def _token_estimate(text: str) -> int:
    return max(1, len(text) // 2)


def _file_type_for_name(file_name: str) -> str:
    return _file_suffix(file_name).lstrip(".") or "unknown"


def _chunk_metadata(
    *,
    draft: ChunkDraft,
    file_uuid: str,
    file_name: str,
    file_type: str,
    sso_user_id: str,
    source_type: str,
    usage_type: str,
    review_status: str,
    rag_scope: str,
    permission_scope: str,
) -> dict:
    base = {
        "document_id": file_uuid,
        "file_name": file_name,
        "file_type": file_type,
        "source_type": usage_type,
        "storage_source_type": source_type,
        "user_id": sso_user_id,
        "section_path": _limit_metadata_text(draft.section_path or draft.section_title),
        "page_or_sheet": _limit_metadata_text(draft.page_or_sheet),
        "chunk_type": draft.chunk_type,
        "chunk_index": draft.chunk_index,
        "usage_type": usage_type,
        "review_status": review_status,
        "rag_scope": rag_scope,
        "permission_scope": permission_scope,
    }
    if draft.metadata:
        base.update(draft.metadata)
    return base


def _metadata_with_embedding(
    *,
    chunk_id: str,
    text: str,
    metadata: dict,
    embedding_service: EmbeddingService | None = None,
) -> tuple[dict, str]:
    service = embedding_service or EmbeddingService()
    vector = service.embed_chunk(text, metadata)
    enriched = {
        **metadata,
        "embedding": service.to_metadata(vector),
    }
    return enriched, service.embedding_id(chunk_id, vector)


def _summary_from_drafts(drafts: list[ChunkDraft]) -> str:
    parts: list[str] = []
    for draft in drafts:
        text = " ".join(draft.text.split())
        if text:
            parts.append(text)
        if sum(len(part) for part in parts) >= 300:
            break
    return " ".join(parts)[:500]


def _refresh_official_knowledge_index(
    db: Session,
    *,
    file_record: KnowledgeFile,
    chunks: list[KnowledgeChunk],
    embedding_service: EmbeddingService | None,
    replace_existing: bool = False,
) -> None:
    from .config import get_settings
    from .knowledge_cache import RedisKnowledgeCache
    from .knowledge_keyword_index import TantivyKnowledgeIndex
    from .knowledge_search import _query_terms, clear_knowledge_search_caches
    from .knowledge_vector_index import QdrantKnowledgeIndex

    clear_knowledge_search_caches()
    settings = get_settings()
    RedisKnowledgeCache.from_settings(settings).bump_knowledge_version()
    is_official = (
        file_record.usage_type == "official_knowledge"
        and file_record.rag_enabled
        and file_record.review_status in {"approved", "official"}
        and file_record.rag_scope == "company"
        and file_record.permission_scope == "company"
    )
    keyword_index = TantivyKnowledgeIndex.from_settings(settings)
    if not is_official:
        keyword_index.delete_file(file_record.uuid)
        return
    keyword_rows: list[tuple[str, str]] = []
    cipher = ContentCipher(settings.content_encryption_key)
    for chunk in chunks:
        payload = cipher.decrypt_json(
            EncryptedPayload(
                ciphertext=chunk.chunk_text_ciphertext,
                nonce=chunk.chunk_text_nonce,
            ),
            chunk.chunk_id.encode(),
        )
        haystack = "\n".join([
            file_record.file_name,
            chunk.section_title,
            str(payload.get("text", "")),
        ])
        keyword_rows.append((chunk.chunk_id, " ".join(_query_terms(haystack))))
    keyword_index.replace_file(file_record.uuid, keyword_rows)
    service = embedding_service or EmbeddingService()
    index = QdrantKnowledgeIndex.from_settings(
        settings,
        dimensions=getattr(service, "dimensions", 0),
    )
    if not index.enabled:
        return
    try:
        if replace_existing:
            index.delete_file(file_record.uuid)
        index.upsert_rows(
            [(chunk, file_record) for chunk in chunks],
            embedding_service=service,
        )
    except Exception:
        # The persisted MySQL vectors remain the authoritative fallback.
        return


def invalidate_knowledge_search(
    *,
    file_uuid: str = "",
    remove_vector_points: bool = False,
) -> None:
    """Invalidate cross-worker retrieval caches after knowledge visibility changes."""
    from .config import get_settings
    from .knowledge_cache import RedisKnowledgeCache
    from .knowledge_keyword_index import TantivyKnowledgeIndex
    from .knowledge_search import clear_knowledge_search_caches
    from .knowledge_vector_index import QdrantKnowledgeIndex

    clear_knowledge_search_caches()
    settings = get_settings()
    RedisKnowledgeCache.from_settings(settings).bump_knowledge_version()
    if not remove_vector_points or not file_uuid:
        return
    TantivyKnowledgeIndex.from_settings(settings).delete_file(file_uuid)
    index = QdrantKnowledgeIndex.from_settings(settings, dimensions=0)
    if not index.enabled:
        return
    try:
        index.delete_file(file_uuid)
    except Exception:
        # MySQL status and permission checks still prevent stale points from leaking.
        return


def hard_delete_session_attachment_files(
    db: Session,
    *,
    conversation_id: str,
    storage_root: str | None,
) -> int:
    """Remove files that are scoped exclusively to a permanently deleted chat."""
    files = list(db.scalars(
        select(KnowledgeFile).where(
            KnowledgeFile.conversation_id == conversation_id,
            KnowledgeFile.usage_type == "session_attachment",
            KnowledgeFile.hard_deleted_at.is_(None),
        )
    ))
    if not files:
        return 0

    root = Path(storage_root).expanduser().resolve() if storage_root else None
    now = datetime.now()
    for file_record in files:
        db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.file_id == file_record.id))
        if root is not None and file_record.file_path:
            stored_path = Path(file_record.file_path).expanduser().resolve()
            if _is_relative_to(stored_path, root) and stored_path.is_file():
                stored_path.unlink()
        file_record.status = "HARD_DELETED"
        file_record.hard_deleted_at = now
        file_record.rag_enabled = False
        file_record.external_public = False
        file_record.external_download_allowed = False
        file_record.reference_enabled = False
        file_record.archived_at = None
        file_record.file_path = ""
        file_record.stored_file_name = ""

    db.flush()
    for file_record in files:
        invalidate_knowledge_search(file_uuid=file_record.uuid, remove_vector_points=True)
    return len(files)


def create_knowledge_file_from_bytes(
    db: Session,
    *,
    sso_user_id: str,
    file_name: str,
    content: bytes,
    content_type: str,
    cipher: ContentCipher,
    key_version: str,
    visibility: str = "PRIVATE",
    source_type: str = "user_upload",
    source_origin: str = "upload",
    web_capture_id: str = "",
    source_url: str = "",
    usage_type: str = "personal_reference",
    review_status: str = "draft",
    rag_enabled: bool = False,
    reference_enabled: bool = True,
    rag_scope: str = "personal",
    permission_scope: str = "private",
    owner_user_id: str | None = None,
    conversation_id: str = "",
    category: str = "个人素材",
    document_type: str = "其他",
    tags: list[str] | None = None,
    uploaded_by: str | None = None,
    knowledge_base_id: int | None = None,
    target_chars: int = 1000,
    max_chars: int = 1500,
    overlap_chars: int = 150,
    storage_root: str | None = None,
    extra_metadata: dict | None = None,
    file_type_override: str | None = None,
    embedding_service: EmbeddingService | None = None,
) -> tuple[KnowledgeFile, list[KnowledgeChunk]]:
    safe_name = _safe_file_name(file_name)
    if len(content) > MAX_KNOWLEDGE_FILE_BYTES:
        raise HTTPException(status_code=413, detail=f"知识文件大小不能超过 {MAX_KNOWLEDGE_FILE_MB} MB")

    parsed_blocks = _extract_blocks(safe_name, content)
    file_uuid = str(uuid_lib.uuid4())
    file_type = file_type_override or _file_type_for_name(safe_name)
    stored_file_name, file_path = _persist_original_file(
        storage_root=storage_root,
        usage_type=usage_type,
        file_uuid=file_uuid,
        file_name=safe_name,
        content=content,
    )
    chunk_drafts = chunk_blocks(
        parsed_blocks,
        target_chars=target_chars,
        max_chars=max_chars,
        overlap_chars=overlap_chars,
    )
    file_record = KnowledgeFile(
        uuid=file_uuid,
        knowledge_base_id=knowledge_base_id,
        sso_user_id=sso_user_id,
        file_name=safe_name,
        original_file_name=safe_name,
        stored_file_name=stored_file_name,
        file_path=file_path,
        file_type=file_type,
        file_size=len(content),
        content_sha256=hashlib.sha256(content).hexdigest(),
        visibility=visibility,
        status="READY",
        error_code="",
        key_version=key_version,
        category=category,
        document_type=document_type,
        tags_json=tags or [],
        summary=_summary_from_drafts(chunk_drafts),
        parse_status="parsed",
        index_status="indexed",
        source_type=source_type,
        source_origin=source_origin,
        web_capture_id=web_capture_id,
        source_url=source_url,
        usage_type=usage_type,
        review_status=review_status,
        rag_enabled=rag_enabled,
        reference_enabled=reference_enabled,
        rag_scope=rag_scope,
        permission_scope=permission_scope,
        owner_user_id=owner_user_id or sso_user_id,
        conversation_id=conversation_id,
        uploaded_by=uploaded_by or sso_user_id,
    )
    db.add(file_record)
    db.flush()

    chunk_records: list[KnowledgeChunk] = []
    for draft in chunk_drafts:
        chunk_id = str(uuid_lib.uuid4())
        encrypted = cipher.encrypt_json({"text": draft.text}, chunk_id.encode())
        metadata, embedding_id = _metadata_with_embedding(
            chunk_id=chunk_id,
            text=draft.text,
            metadata=_chunk_metadata(
                draft=draft,
                file_uuid=file_uuid,
                file_name=safe_name,
                file_type=file_type,
                sso_user_id=sso_user_id,
                source_type=source_type,
                usage_type=usage_type,
                review_status=review_status,
                rag_scope=rag_scope,
                permission_scope=permission_scope,
            ) | (extra_metadata or {}),
            embedding_service=embedding_service,
        )
        chunk = KnowledgeChunk(
            chunk_id=chunk_id,
            file_id=file_record.id,
            knowledge_base_id=knowledge_base_id,
            file_name=safe_name,
            chunk_text_ciphertext=encrypted.ciphertext,
            chunk_text_nonce=encrypted.nonce,
            page_number=draft.page_number,
            section_title=_limit_metadata_text(draft.section_title),
            chunk_index=draft.chunk_index,
            token_estimate=_token_estimate(draft.text),
            token_count=_token_estimate(draft.text),
            embedding_id=embedding_id,
            metadata_json=metadata,
            status="READY",
        )
        db.add(chunk)
        chunk_records.append(chunk)
    db.flush()
    _refresh_official_knowledge_index(
        db,
        file_record=file_record,
        chunks=chunk_records,
        embedding_service=embedding_service,
    )
    return file_record, chunk_records


def reparse_knowledge_file_from_existing_chunks(
    db: Session,
    *,
    file_record: KnowledgeFile,
    cipher: ContentCipher,
    target_chars: int = 1000,
    max_chars: int = 1500,
    overlap_chars: int = 150,
    storage_root: str | None = None,
    embedding_service: EmbeddingService | None = None,
) -> list[KnowledgeChunk]:
    source_text = ""
    source_blocks: list[ParsedBlock] = []
    if storage_root and file_record.file_path:
        root = Path(storage_root).expanduser().resolve()
        original_path = Path(file_record.file_path).expanduser().resolve()
        if _is_relative_to(original_path, root) and original_path.is_file():
            source_name = file_record.original_file_name or file_record.file_name
            source_blocks = _extract_blocks(source_name, original_path.read_bytes())
            source_text = "\n\n".join(block.text for block in source_blocks).strip()

    existing_chunks = list(
        db.scalars(
            select(KnowledgeChunk)
            .where(
                KnowledgeChunk.file_id == file_record.id,
                KnowledgeChunk.status == "READY",
                KnowledgeChunk.deleted_at.is_(None),
            )
            .order_by(KnowledgeChunk.chunk_index.asc())
        )
    )
    if not source_text and not existing_chunks:
        raise HTTPException(status_code=422, detail="当前文档没有可重解析的内容")

    if not source_text:
        text_parts: list[str] = []
        for chunk in existing_chunks:
            payload = cipher.decrypt_json(
                EncryptedPayload(
                    ciphertext=chunk.chunk_text_ciphertext,
                    nonce=chunk.chunk_text_nonce,
                ),
                chunk.chunk_id.encode(),
            )
            text = str(payload.get("text", "")).strip()
            if text:
                text_parts.append(text)
        source_text = "\n".join(text_parts).strip()
    if not source_text:
        raise HTTPException(status_code=422, detail="当前文档没有可重解析的文本")

    db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.file_id == file_record.id))
    db.flush()

    rebuilt_chunks: list[KnowledgeChunk] = []
    rebuilt_blocks = source_blocks or _parse_text_blocks(file_record.file_name, source_text.encode("utf-8"))
    for draft in chunk_blocks(
        rebuilt_blocks,
        target_chars=target_chars,
        max_chars=max_chars,
        overlap_chars=overlap_chars,
    ):
        chunk_id = str(uuid_lib.uuid4())
        encrypted = cipher.encrypt_json({"text": draft.text}, chunk_id.encode())
        token_count = _token_estimate(draft.text)
        metadata, embedding_id = _metadata_with_embedding(
            chunk_id=chunk_id,
            text=draft.text,
            metadata=_chunk_metadata(
                draft=draft,
                file_uuid=file_record.uuid,
                file_name=file_record.file_name,
                file_type=_file_type_for_name(file_record.file_name),
                sso_user_id=file_record.sso_user_id,
                source_type=file_record.source_type,
                usage_type=file_record.usage_type,
                review_status=file_record.review_status,
                rag_scope=file_record.rag_scope,
                permission_scope=file_record.permission_scope,
            ),
            embedding_service=embedding_service,
        )
        chunk = KnowledgeChunk(
            chunk_id=chunk_id,
            file_id=file_record.id,
            knowledge_base_id=file_record.knowledge_base_id,
            file_name=file_record.file_name,
            chunk_text_ciphertext=encrypted.ciphertext,
            chunk_text_nonce=encrypted.nonce,
            page_number=draft.page_number,
            section_title=_limit_metadata_text(draft.section_title),
            chunk_index=draft.chunk_index,
            token_estimate=token_count,
            token_count=token_count,
            embedding_id=embedding_id,
            metadata_json=metadata,
            status="READY",
        )
        db.add(chunk)
        rebuilt_chunks.append(chunk)
    file_record.parse_status = "parsed"
    file_record.index_status = "indexed"
    file_record.error_code = ""
    db.flush()
    _refresh_official_knowledge_index(
        db,
        file_record=file_record,
        chunks=rebuilt_chunks,
        embedding_service=embedding_service,
        replace_existing=True,
    )
    return rebuilt_chunks

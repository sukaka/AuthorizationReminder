import asyncio
import os
from io import BytesIO
from zipfile import ZipFile

import pytest
from docx import Document
from sqlalchemy import select
from starlette.datastructures import Headers


def _upload(
    client,
    *,
    task_uuid: str,
    file_name: str = "meeting.txt",
    content: bytes = "会议内容".encode("utf-8"),
    content_type: str = "text/plain",
):
    return client.post(
        "/api/ai/attachments",
        data={"task_uuid": task_uuid},
        files={"file": (file_name, content, content_type)},
    )


def _build_docx_bytes(*, paragraph_text: str = "会议段落") -> bytes:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph(paragraph_text)
    document.save(buffer)
    return buffer.getvalue()


def _build_docx_table_bytes() -> bytes:
    buffer = BytesIO()
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "事项"
    table.cell(0, 1).text = "负责人"
    table.cell(1, 0).text = "项目甲"
    document.save(buffer)
    return buffer.getvalue()


def _build_zip_bytes(entries: dict[str, bytes]) -> bytes:
    buffer = BytesIO()
    with ZipFile(buffer, "w") as archive:
        for name, content in entries.items():
            archive.writestr(name, content)
    return buffer.getvalue()


def _build_docx_package_bytes(document_xml: bytes) -> bytes:
    return _build_zip_bytes({
        "[Content_Types].xml": b"""<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>""",
        "_rels/.rels": b"""<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>""",
        "word/document.xml": document_xml,
    })


def _load_attachment_text(generation_db, seeded_task, attachment_uuid: str) -> str:
    from app.attachments import load_owned_attachment_texts
    from app.crypto import ContentCipher

    cipher = ContentCipher(os.environ["CONTENT_ENCRYPTION_KEY"])
    _, extracted_text = load_owned_attachment_texts(
        generation_db,
        "dev",
        seeded_task.id,
        [attachment_uuid],
        cipher,
    )[0]
    return extracted_text


def test_upload_txt_attachment_extracts_and_encrypts_text(
    generation_client,
    generation_db,
    seeded_task,
):
    from app.models import GenerationAttachment

    response = _upload(generation_client, task_uuid=seeded_task.uuid)

    assert response.status_code == 201
    body = response.json()
    assert body["uuid"]
    assert body["name"] == "meeting.txt"
    assert body["type"] == "text/plain"
    assert body["size"] == len("会议内容".encode("utf-8"))
    assert body["created_at"]
    assert body["status"] == "READY"
    assert body["extracted_characters"] == 4

    attachment = generation_db.scalar(
        select(GenerationAttachment).where(GenerationAttachment.uuid == body["uuid"])
    )
    assert attachment is not None
    assert attachment.generation_id is None
    assert attachment.file_name == "meeting.txt"
    assert attachment.extracted_text_ciphertext
    assert "会议内容".encode("utf-8") not in attachment.extracted_text_ciphertext


def test_upload_md_attachment_is_supported(
    generation_client,
    seeded_task,
):
    response = _upload(
        generation_client,
        task_uuid=seeded_task.uuid,
        file_name="notes.md",
        content="# 纪要".encode("utf-8"),
        content_type="text/markdown",
    )

    assert response.status_code == 201
    body = response.json()
    assert body["name"] == "notes.md"
    assert body["type"] == "text/markdown"
    assert body["extracted_characters"] == 4


def test_upload_docx_attachment_extracts_paragraph_text(
    generation_client,
    generation_db,
    seeded_task,
):
    text = "会议段落"

    response = _upload(
        generation_client,
        task_uuid=seeded_task.uuid,
        file_name="meeting.docx",
        content=_build_docx_bytes(paragraph_text=text),
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )

    assert response.status_code == 201
    body = response.json()
    assert body["name"] == "meeting.docx"
    assert body["extracted_characters"] >= len(text)
    assert text in _load_attachment_text(generation_db, seeded_task, body["uuid"])


def test_upload_docx_attachment_extracts_table_text(
    generation_client,
    generation_db,
    seeded_task,
):
    response = _upload(
        generation_client,
        task_uuid=seeded_task.uuid,
        file_name="table.docx",
        content=_build_docx_table_bytes(),
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )

    assert response.status_code == 201

    extracted_text = _load_attachment_text(
        generation_db,
        seeded_task,
        response.json()["uuid"],
    )
    assert "事项" in extracted_text
    assert "负责人" in extracted_text
    assert "项目甲" in extracted_text
    assert "待确认" in extracted_text


def test_upload_bad_docx_attachment_returns_clear_parse_error(
    generation_client,
    seeded_task,
):
    response = _upload(
        generation_client,
        task_uuid=seeded_task.uuid,
        file_name="bad.docx",
        content=b"not a docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )

    assert response.status_code == 422
    assert "DOCX" in response.text
    assert "无法解析" in response.text


def test_upload_docx_with_invalid_document_body_returns_clear_parse_error(
    generation_client,
    seeded_task,
):
    response = _upload(
        generation_client,
        task_uuid=seeded_task.uuid,
        file_name="invalid-body.docx",
        content=_build_docx_package_bytes(
            b"""<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:notBody/>
</w:document>"""
        ),
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )

    assert response.status_code == 422
    assert "DOCX" in response.text
    assert "无法解析" in response.text


@pytest.mark.parametrize(
    "content",
    [
        _build_zip_bytes({}),
        _build_zip_bytes({"[Content_Types].xml": b"<Types>"}),
    ],
    ids=["empty-zip", "malformed-content-types"],
)
def test_upload_malformed_zip_docx_attachment_returns_clear_parse_error(
    generation_client,
    seeded_task,
    content,
):
    response = _upload(
        generation_client,
        task_uuid=seeded_task.uuid,
        file_name="malformed.docx",
        content=content,
        content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )

    assert response.status_code == 422
    assert "DOCX" in response.text
    assert "无法解析" in response.text


def test_upload_attachment_requires_ai_assistant_use_permission(
    monkeypatch,
    generation_client,
    seeded_task,
):
    from app import main as main_module

    calls: list[str] = []

    async def fake_require_action(action, request, session, settings, *, resource=None):
        calls.append(action)
        return session

    monkeypatch.setattr(main_module, "require_action", fake_require_action)

    response = _upload(generation_client, task_uuid=seeded_task.uuid)

    assert response.status_code == 201
    assert calls == ["ai_assistant:use"]


def test_create_attachment_does_not_commit(
    monkeypatch,
    generation_db,
    seeded_task,
):
    from app.attachments import create_attachment
    from app.crypto import ContentCipher
    from fastapi import UploadFile

    def fail_commit():
        raise AssertionError("create_attachment must not commit")

    monkeypatch.setattr(generation_db, "commit", fail_commit)
    upload = UploadFile(
        BytesIO("会议内容".encode("utf-8")),
        filename="meeting.txt",
        headers=Headers({"content-type": "text/plain"}),
    )
    cipher = ContentCipher(os.environ["CONTENT_ENCRYPTION_KEY"])

    attachment, extracted_characters = asyncio.run(
        create_attachment(
            generation_db,
            "dev",
            seeded_task.uuid,
            upload,
            cipher,
            "v1",
        )
    )

    assert attachment.id is not None
    assert attachment.generation_id is None
    assert extracted_characters == 4


def test_upload_rolls_back_attachment_when_audit_fails(
    monkeypatch,
    generation_client,
    generation_db,
    seeded_task,
):
    from app import main as main_module
    from app.models import GenerationAttachment

    def fail_audit(*args, **kwargs):
        raise RuntimeError("audit failed")

    monkeypatch.setattr(main_module, "write_request_audit", fail_audit)

    with pytest.raises(RuntimeError, match="audit failed"):
        _upload(generation_client, task_uuid=seeded_task.uuid)
    assert generation_db.scalar(select(GenerationAttachment)) is None


def test_upload_unsupported_attachment_type_returns_clear_error(
    generation_client,
    seeded_task,
):
    response = _upload(
        generation_client,
        task_uuid=seeded_task.uuid,
        file_name="image.png",
        content=b"content",
        content_type="image/png",
    )

    assert response.status_code == 415
    assert response.json()["detail"] == "当前仅支持 txt、md、docx"


def test_upload_pdf_attachment_returns_clear_not_yet_supported_error(
    generation_client,
    seeded_task,
):
    response = _upload(
        generation_client,
        task_uuid=seeded_task.uuid,
        file_name="report.pdf",
        content=b"%PDF-1.7",
        content_type="application/pdf",
    )

    assert response.status_code == 422
    assert response.json()["detail"] == "PDF 文本提取将在下一步启用；扫描件暂不支持 OCR"


def test_upload_attachment_rejects_non_utf8_text(
    generation_client,
    seeded_task,
):
    response = _upload(
        generation_client,
        task_uuid=seeded_task.uuid,
        content=b"\xff\xfe",
    )

    assert response.status_code == 422
    assert response.json()["detail"] == "文本附件必须使用 UTF-8 编码"


def test_upload_attachment_rejects_files_larger_than_20mb(
    generation_client,
    seeded_task,
):
    response = _upload(
        generation_client,
        task_uuid=seeded_task.uuid,
        content=b"x" * (20 * 1024 * 1024 + 1),
    )

    assert response.status_code == 413
    assert response.json()["detail"] == "附件大小不能超过 20 MB"


def test_upload_attachment_rejects_file_name_longer_than_255(
    generation_client,
    seeded_task,
):
    response = _upload(
        generation_client,
        task_uuid=seeded_task.uuid,
        file_name=f"{'a' * 256}.txt",
    )

    assert response.status_code == 422
    assert response.json()["detail"] == "文件名不能超过 255 个字符"

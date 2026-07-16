from io import BytesIO

from docx import Document

from app.chat_word_export import TemplateRenderer
from app.professional_delivery.export_service import (
    structured_content_export_report,
    structured_content_to_markdown,
)


def test_structured_content_round_trips_through_office_renderer() -> None:
    content = {
        "schema_version": "2",
        "blocks": [
            {"block_id": "heading", "type": "heading", "level": 2, "text": "季度风险结论"},
            {"block_id": "summary", "type": "paragraph", "text": "本季度高风险事件已全部闭环。"},
            {
                "block_id": "risk-table",
                "type": "table",
                "rows": [["风险等级", "数量"], ["高", "3"]],
            },
            {"block_id": "architecture-image", "type": "image", "asset_id": "asset-1", "alt": "系统截图"},
        ],
    }

    markdown = structured_content_to_markdown(content, fallback="")
    assert "## 季度风险结论" in markdown
    assert "风险等级" in markdown
    assert "图片：系统截图" in markdown
    assert structured_content_to_markdown(content, fallback="") == markdown

    rendered = TemplateRenderer().render(
        title="季度风险报告",
        task_name="安全运营",
        department="安全部",
        author="测试用户",
        output=markdown,
        version="V1",
        template_name="juxin_standard",
    )
    document = Document(BytesIO(rendered))
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    office_text = f"{paragraph_text}\n{table_text}"
    assert "季度风险结论" in office_text
    assert "本季度高风险事件已全部闭环。" in office_text
    assert "风险等级" in office_text
    assert "图片：系统截图" in office_text


def test_structured_content_export_reports_degraded_and_rejected_blocks() -> None:
    content = {
        "schema_version": "2",
        "blocks": [
            {"block_id": "paragraph", "type": "paragraph", "text": "正文"},
            {"block_id": "media", "type": "image", "asset_id": "asset-1", "alt": "图"},
            {"block_id": "future", "type": "smart_chart", "text": "不可丢失"},
            {"block_id": "divider", "type": "divider"},
        ],
    }

    report = structured_content_export_report(content)

    assert report["status"] == "degraded"
    assert report["supported_features"] == ["divider", "image", "paragraph"]
    assert [item["code"] for item in report["degraded_features"]] == [
        "media_asset_not_embedded",
    ]
    assert report["rejected_features"] == [
        {
            "code": "unsupported_block_type",
            "message": "内容块类型未被 DOCX 渲染器支持，已保留显式占位",
            "block_index": 2,
            "block_type": "smart_chart",
        }
    ]
    markdown = structured_content_to_markdown(content, fallback="")
    assert "【未支持内容块：smart_chart】：不可丢失" in markdown
    assert "---" in markdown


def test_structured_content_export_marks_known_blocks_that_cannot_be_rendered() -> None:
    content = {
        "blocks": [
            {"type": "paragraph"},
            {"type": "table", "rows": [[""]]},
            {"type": "image", "asset_id": "asset-1"},
        ]
    }

    report = structured_content_export_report(content)
    assert report["status"] == "degraded"
    assert [item["code"] for item in report["degraded_features"]] == [
        "empty_block",
        "empty_table",
        "media_asset_not_embedded",
    ]
    markdown = structured_content_to_markdown(content, fallback="")
    assert "【空内容块：第 1 个块】" in markdown
    assert "【表格无法导出：第 2 个块】" in markdown
    assert "【图片未嵌入：第 3 个块】" in markdown

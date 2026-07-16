import json
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from docx.shared import Inches
from fastapi import HTTPException


DOCX_IMPORT_REPORT_GOLDEN = Path(__file__).parent / "fixtures" / "docx_import_report_golden.json"


def _docx_bytes(*, include_image: bool = False) -> bytes:
    document = Document()
    document.add_heading("安全运营月报", level=1)
    document.add_paragraph("本月完成全部安全巡检。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "结果"
    table.cell(1, 0).text = "重大事件"
    table.cell(1, 1).text = "0"
    if include_image:
        document.add_picture(BytesIO(
            b"\x89PNG\r\n\x1a\n"
            b"\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
            b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89"
            b"\x00\x00\x00\x0dIDAT\x08\xd7c\xfc\xcf\xc0\xf0\x1f\x00\x05\x00\x01\xff\x89\x99=\x1d"
            b"\x00\x00\x00\x00IEND\xaeB`\x82"
        ), width=Inches(1))
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _rewrite_docx_package(
    data: bytes,
    *,
    document_replacements: tuple[tuple[bytes, bytes], ...] = (),
    extra_entries: dict[str, bytes] | None = None,
) -> bytes:
    """Create a deterministic DOCX fixture with one targeted OOXML feature."""
    extras = extra_entries or {}
    output = BytesIO()
    with ZipFile(BytesIO(data)) as source, ZipFile(output, "w") as target:
        for info in source.infolist():
            payload = source.read(info.filename)
            if info.filename == "word/document.xml":
                for before, after in document_replacements:
                    payload = payload.replace(before, after)
            target.writestr(info, payload)
        for name, payload in extras.items():
            target.writestr(name, payload)
    return output.getvalue()


def test_docx_import_preserves_body_order_and_table_structure() -> None:
    from app.professional_delivery.docx_import import structured_content_from_docx

    result = structured_content_from_docx(_docx_bytes())

    assert result["schema_version"] == "2"
    assert [block["type"] for block in result["blocks"]] == ["heading", "paragraph", "table"]
    assert result["blocks"][0]["text"] == "安全运营月报"
    assert result["blocks"][1]["text"] == "本月完成全部安全巡检。"
    assert result["blocks"][2]["rows"] == [
        {"cells": [{"text": "指标"}, {"text": "结果"}]},
        {"cells": [{"text": "重大事件"}, {"text": "0"}]},
    ]
    assert all(block["block_id"].startswith("docx-") for block in result["blocks"])
    assert result["import_meta"] == {
        "source_format": "docx",
        "warnings": [],
        "media_count": 0,
        "import_report": {
            "status": "supported",
            "supported_features": ["heading", "paragraph", "table"],
            "degraded_features": [],
            "rejected_features": [],
        },
    }


def test_docx_import_reports_media_without_fabricating_asset_reference() -> None:
    from app.professional_delivery.docx_import import structured_content_from_docx

    result = structured_content_from_docx(_docx_bytes(include_image=True))

    assert result["import_meta"]["media_count"] == 1
    assert result["import_meta"]["warnings"] == ["media_not_imported"]
    assert result["import_meta"]["import_report"] == {
        "status": "degraded",
        "supported_features": ["heading", "paragraph", "table"],
        "degraded_features": [
            {"code": "media_not_imported", "message": "DOCX 内容已降级处理"}
        ],
        "rejected_features": [],
    }
    assert all(block["type"] != "media" for block in result["blocks"])


def test_docx_import_extracts_related_image_and_preserves_block_order() -> None:
    from app.professional_delivery.docx_import import (
        structured_content_and_media_from_docx,
    )

    result, media = structured_content_and_media_from_docx(_docx_bytes(include_image=True))

    assert [block["type"] for block in result["blocks"]] == [
        "heading",
        "paragraph",
        "table",
        "image",
    ]
    image_block = result["blocks"][-1]
    assert image_block["source_id"] == media[0].source_id
    assert image_block["mime_type"] == "image/png"
    assert image_block["size_bytes"] == len(media[0].data)
    assert media[0].original_file_name == "image1.png"
    assert media[0].data.startswith(b"\x89PNG\r\n\x1a\n")


def test_docx_import_rejects_invalid_archive_with_domain_error() -> None:
    from app.professional_delivery.docx_import import structured_content_from_docx

    with pytest.raises(HTTPException) as exc_info:
        structured_content_from_docx(b"not-a-docx")

    assert exc_info.value.status_code == 422
    assert exc_info.value.detail == {
        "code": "INVALID_DOCX_IMPORT",
        "message": "DOCX 文件无法解析",
    }


def test_docx_import_reports_complex_word_features_without_silent_loss() -> None:
    from app.professional_delivery.docx_import import structured_content_from_docx

    source = _docx_bytes()
    fixture = _rewrite_docx_package(
        source,
        document_replacements=(
            (b"</w:pPr>", b'<w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr></w:pPr>'),
            (b"</w:body>", b'<w:altChunk r:id="rId999"/></w:body>'),
        ),
        extra_entries={
            "word/header1.xml": b"<w:hdr xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\"/>",
            "word/vbaProject.bin": b"not executable in tests",
        },
    )

    result = structured_content_from_docx(fixture)
    report = result["import_meta"]["import_report"]
    assert report["status"] == "degraded"
    assert [item["code"] for item in report["degraded_features"]] == [
        "headers_footers",
        "lists",
    ]
    assert [item["code"] for item in report["rejected_features"]] == [
        "macros",
        "alt_chunk",
    ]
    assert result["import_meta"]["warnings"] == [
        "headers_footers",
        "macros",
        "lists",
        "alt_chunk",
    ]


def test_docx_import_reports_comments_and_fields_deterministically() -> None:
    from app.professional_delivery.docx_import import structured_content_from_docx

    fixture = _rewrite_docx_package(
        _docx_bytes(),
        document_replacements=(
            (
                b"</w:body>",
                b"<w:p><w:r><w:commentReference w:id=\"0\"/></w:r></w:p>"
                b"<w:p><w:r><w:fldChar w:fldCharType=\"begin\"/>"
                b"<w:instrText> PAGE </w:instrText>"
                b"<w:fldChar w:fldCharType=\"end\"/></w:r></w:p></w:body>",
            ),
        ),
        extra_entries={
            "word/comments.xml": b"<w:comments xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\"/>"
        },
    )

    result = structured_content_from_docx(fixture)
    report = result["import_meta"]["import_report"]
    assert report["status"] == "degraded"
    assert [item["code"] for item in report["degraded_features"]] == [
        "comments",
        "fields",
    ]
    assert report["rejected_features"] == []
    assert result["import_meta"]["warnings"] == ["comments", "fields"]


def test_docx_import_reports_floating_images_as_degraded() -> None:
    from app.professional_delivery.docx_import import structured_content_and_media_from_docx

    source = _docx_bytes(include_image=True)
    fixture = _rewrite_docx_package(
        source,
        document_replacements=(
            # The inline element carries namespace attributes, so replace its
            # qualified-name prefix rather than assuming an empty start tag.
            (b"<wp:inline", b"<wp:anchor"),
            (b"</wp:inline>", b"</wp:anchor>"),
        ),
    )
    content, media = structured_content_and_media_from_docx(fixture)
    report = content["import_meta"]["import_report"]
    assert len(media) == 1
    assert report["status"] == "degraded"
    assert [item["code"] for item in report["degraded_features"]] == ["floating_layout"]
    assert report["rejected_features"] == []


def test_docx_import_report_matches_complex_feature_golden() -> None:
    """Keep the editor's supported/degraded/rejected contract reviewable.

    The fixtures are generated in memory so this test does not need to check
    in binary DOCX files; the JSON golden captures the externally visible
    report that clients use before opening an imported document in the editor.
    """
    from app.professional_delivery.docx_import import (
        structured_content_and_media_from_docx,
        structured_content_from_docx,
    )

    source = _docx_bytes()
    complex_fixture = _rewrite_docx_package(
        source,
        document_replacements=(
            (
                b"</w:pPr>",
                b'<w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr></w:pPr>',
            ),
            (b"</w:body>", b'<w:altChunk r:id="rId999"/></w:body>'),
        ),
        extra_entries={
            "word/header1.xml": b"<w:hdr xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\"/>",
            "word/vbaProject.bin": b"not executable in tests",
        },
    )
    comments_fields_fixture = _rewrite_docx_package(
        source,
        document_replacements=(
            (
                b"</w:body>",
                b"<w:p><w:r><w:commentReference w:id=\"0\"/></w:r></w:p>"
                b"<w:p><w:r><w:fldChar w:fldCharType=\"begin\"/>"
                b"<w:instrText> PAGE </w:instrText>"
                b"<w:fldChar w:fldCharType=\"end\"/></w:r></w:p></w:body>",
            ),
        ),
        extra_entries={
            "word/comments.xml": b"<w:comments xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\"/>"
        },
    )
    floating_fixture = _rewrite_docx_package(
        _docx_bytes(include_image=True),
        document_replacements=(
            (b"<wp:inline", b"<wp:anchor"),
            (b"</wp:inline>", b"</wp:anchor>"),
        ),
    )
    cases = {
        "basic": structured_content_from_docx(source),
        "media_compat": structured_content_from_docx(_docx_bytes(include_image=True)),
        "complex": structured_content_from_docx(complex_fixture),
        "comments_fields": structured_content_from_docx(comments_fields_fixture),
        "floating_image": structured_content_and_media_from_docx(floating_fixture)[0],
    }
    golden = json.loads(DOCX_IMPORT_REPORT_GOLDEN.read_text(encoding="utf-8"))
    assert set(cases) == set(golden)
    for name, content in cases.items():
        assert content["import_meta"]["import_report"] == golden[name], name

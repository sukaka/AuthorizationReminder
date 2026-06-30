from io import BytesIO

from docx import Document


def _encrypt(cipher, message_uuid: str, content: str):
    payload = cipher.encrypt_json({"content": content}, message_uuid.encode())
    return payload.ciphertext, payload.nonce


def _seed_chat(generation_db, *, user_id: str = "u-1"):
    from app.config import get_settings
    from app.crypto import ContentCipher
    from app.models import ChatMessage, ChatMessageSource, ChatSession

    cipher = ContentCipher(get_settings().content_encryption_key)
    session = ChatSession(
        uuid="chat-export-session",
        sso_user_id=user_id,
        title="导出测试会话",
        mode="NORMAL",
        status="active",
    )
    generation_db.add(session)
    generation_db.flush()

    user_ciphertext, user_nonce = _encrypt(cipher, "chat-user-message", "请整理交付方案")
    answer_ciphertext, answer_nonce = _encrypt(
        cipher,
        "chat-assistant-message",
        "# 项目交付方案\n\n- 部署准备\n- 用户培训\n\n| 阶段 | 内容 |\n|---|---|\n| 验收 | 待确认 |",
    )
    user_message = ChatMessage(
            uuid="chat-user-message",
            session_id=session.id,
            sso_user_id=user_id,
            role="user",
            content_ciphertext=user_ciphertext,
            content_nonce=user_nonce,
            key_version="v1",
            status="COMPLETED",
    )
    assistant_message = ChatMessage(
            uuid="chat-assistant-message",
            session_id=session.id,
            sso_user_id=user_id,
            role="assistant",
            content_ciphertext=answer_ciphertext,
            content_nonce=answer_nonce,
            key_version="v1",
            status="COMPLETED",
    )
    generation_db.add_all([user_message, assistant_message])
    generation_db.flush()
    generation_db.add_all([
        ChatMessageSource(
            message_id=assistant_message.id,
            source_type="official_knowledge",
            source_uuid="official-file",
            title="正式白皮书",
            file_name="聚信产品白皮书.pdf",
            chunk_id="official-chunk-001",
            page_number=12,
            section_title="部署方式",
            chunk_index=0,
            score=9,
        ),
        ChatMessageSource(
            message_id=assistant_message.id,
            source_type="official_knowledge",
            source_uuid="official-file-duplicate",
            title="正式白皮书重复命中",
            file_name="聚信产品白皮书.pdf",
            chunk_id="official-chunk-duplicate",
            page_number=12,
            section_title="部署方式",
            chunk_index=1,
            score=8,
        ),
        ChatMessageSource(
            message_id=assistant_message.id,
            source_type="personal_reference",
            source_uuid="personal-file",
            title="我的会议记录",
            file_name="我的会议记录.docx",
            chunk_id="personal-chunk-001",
            page_number=None,
            section_title="会议讨论内容",
            chunk_index=0,
            score=5,
        ),
        ChatMessageSource(
            message_id=assistant_message.id,
            source_type="session_attachment",
            source_uuid="session-file",
            title="当前会话附件",
            file_name="客户访谈记录.pdf",
            chunk_id="session-chunk-001",
            page_number=3,
            section_title="客户诉求",
            chunk_index=0,
            score=7,
        ),
    ])
    generation_db.commit()
    return session


def test_chat_single_answer_word_export_creates_downloadable_docx(
    client_for_user,
    generation_db,
    tmp_path,
):
    from app.config import get_settings
    from app.main import app
    from app.models import ExportRecord

    session = _seed_chat(generation_db)
    settings = get_settings().model_copy(update={"export_storage_dir": str(tmp_path)})
    app.dependency_overrides[get_settings] = lambda: settings
    try:
        client = client_for_user("u-1")

        response = client.post(
            "/api/export/word",
            json={
                "conversation_id": session.uuid,
                "message_id": "chat-assistant-message",
                "export_type": "single_answer",
                "template": "juxin_standard",
                "format_before_export": False,
            },
        )

        assert response.status_code == 201
        payload = response.json()
        assert payload["file_name"].endswith(".docx")
        assert payload["download_url"].startswith("/api/export/download/")
        assert "file_path" not in payload

        download = client.get(payload["download_url"])
        assert download.status_code == 200
        assert download.headers["content-type"].startswith(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        document = Document(BytesIO(download.content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        assert "聚信得仁" in document.sections[0].header.paragraphs[0].text
        assert "项目交付方案" in text
        assert "部署准备" in text
        assert "参考来源" in text
        assert "聚信产品白皮书.pdf——公司知识库 / 正式知识来源，第 12 页，部署方式" in text
        assert "我的会议记录.docx——我的上传文件，仅用于本次内容生成，会议讨论内容" in text
        assert "客户访谈记录.pdf——当前会话附件，第 3 页，客户诉求" in text
        assert "本文参考用户个人上传资料生成，仅供用户本人使用" in text
        assert "本文参考当前会话附件生成，仅供本次会话使用" in text
        assert "official-chunk-001" not in text
        assert "official-chunk-duplicate" not in text
        assert "personal-chunk-001" not in text
        assert "session-chunk-001" not in text
        assert any(table.cell(0, 0).text == "阶段" for table in document.tables)

        record = generation_db.query(ExportRecord).one()
        assert record.conversation_id == session.uuid
        assert record.message_id == "chat-assistant-message"
        assert record.export_type == "single_answer"
        assert record.created_by == "u-1"
    finally:
        app.dependency_overrides.pop(get_settings, None)


def test_reference_sources_markdown_deduplicates_and_keeps_source_notices(
    generation_db,
):
    from app.chat_word_export import _reference_sources_markdown
    from app.models import ChatMessage

    _seed_chat(generation_db)
    assistant_message = generation_db.query(ChatMessage).filter_by(
        uuid="chat-assistant-message",
    ).one()

    markdown = _reference_sources_markdown(generation_db, [assistant_message])

    assert "1. 聚信产品白皮书.pdf——公司知识库 / 正式知识来源，第 12 页，部署方式" in markdown
    assert "2. 我的会议记录.docx——我的上传文件，仅用于本次内容生成，会议讨论内容" in markdown
    assert "3. 客户访谈记录.pdf——当前会话附件，第 3 页，客户诉求" in markdown
    assert "4. " not in markdown
    assert "official-chunk-duplicate" not in markdown
    assert "本文参考用户个人上传资料生成，仅供用户本人使用。" in markdown
    assert "本文参考当前会话附件生成，仅供本次会话使用。" in markdown


def test_formal_document_word_export_respects_message_id_scope(
    client_for_user,
    generation_db,
    tmp_path,
):
    from app.config import get_settings
    from app.main import app
    from app.models import ExportRecord

    session = _seed_chat(generation_db)
    settings = get_settings().model_copy(update={"export_storage_dir": str(tmp_path)})
    app.dependency_overrides[get_settings] = lambda: settings
    try:
        client = client_for_user("u-1")

        response = client.post(
            "/api/export/word",
            json={
                "conversation_id": session.uuid,
                "message_id": "chat-assistant-message",
                "export_type": "formal_document",
                "template": "juxin_standard",
                "format_before_export": True,
            },
        )

        assert response.status_code == 201
        download = client.get(response.json()["download_url"])
        assert download.status_code == 200
        document = Document(BytesIO(download.content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        assert "项目交付方案" in text
        assert "请整理交付方案" not in text
        assert "参考来源" in text

        record = generation_db.query(ExportRecord).one()
        assert record.message_id == "chat-assistant-message"
        assert record.export_type == "formal_document"
    finally:
        app.dependency_overrides.pop(get_settings, None)


def test_transient_knowledge_result_word_export_keeps_reference_sources(
    client_for_user,
    generation_db,
    tmp_path,
):
    from app.config import get_settings
    from app.main import app
    from app.models import ExportRecord

    settings = get_settings().model_copy(update={"export_storage_dir": str(tmp_path)})
    app.dependency_overrides[get_settings] = lambda: settings
    try:
        client = client_for_user("u-1")

        response = client.post(
            "/api/export/word/content",
            json={
                "title": "会议纪要模板-文档问答结果",
                "content": "文档回答：验收材料需要包含会议结论、责任人和下一步计划。",
                "template": "juxin_standard",
                "sources": [
                    {
                        "source_kind": "personal_reference",
                        "file_id": "file-personal-1",
                        "file_name": "会议纪要模板.docx",
                        "page_number": 2,
                        "section_title": "验收材料",
                        "chunk_id": "chunk-ask-secret",
                        "score": 90,
                        "snippet": "验收材料包含会议结论、责任人和下一步计划。",
                    }
                ],
            },
        )

        assert response.status_code == 201
        payload = response.json()
        assert payload["download_url"].startswith("/api/export/download/")

        download = client.get(payload["download_url"])
        assert download.status_code == 200
        document = Document(BytesIO(download.content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        assert "会议纪要模板-文档问答结果" in text
        assert "验收材料需要包含会议结论" in text
        assert "参考来源" in text
        assert "会议纪要模板.docx——我的上传文件，仅用于本次内容生成，第 2 页，验收材料" in text
        assert "本文参考用户个人上传资料生成，仅供用户本人使用" in text
        assert "chunk-ask-secret" not in text

        record = generation_db.query(ExportRecord).one()
        assert record.conversation_id == ""
        assert record.export_type == "knowledge_result"
        assert record.created_by == "u-1"
    finally:
        app.dependency_overrides.pop(get_settings, None)


def test_chat_full_conversation_word_export_requires_session_owner(
    client_for_user,
    generation_db,
    tmp_path,
):
    from app.config import get_settings
    from app.main import app

    session = _seed_chat(generation_db, user_id="u-owner")
    settings = get_settings().model_copy(update={"export_storage_dir": str(tmp_path)})
    app.dependency_overrides[get_settings] = lambda: settings
    try:
        client = client_for_user("u-other")

        response = client.post(
            "/api/export/word",
            json={
                "conversation_id": session.uuid,
                "export_type": "full_conversation",
                "template": "juxin_standard",
                "format_before_export": False,
            },
        )

        assert response.status_code == 404
    finally:
        app.dependency_overrides.pop(get_settings, None)

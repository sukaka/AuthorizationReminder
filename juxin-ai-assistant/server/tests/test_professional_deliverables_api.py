from datetime import datetime
from io import BytesIO
from types import SimpleNamespace

import pytest
from docx import Document
from sqlalchemy import func, select


CONTENT = {
    "schema_version": "1",
    "blocks": [
        {
            "block_id": "monthly-overview",
            "type": "paragraph",
            "text": "本月未发生重大安全事件。",
        }
    ],
}

UPDATED_CONTENT = {
    "schema_version": "1",
    "blocks": [
        {
            "block_id": "monthly-overview",
            "type": "paragraph",
            "text": "本月完成全部安全巡检，未发生重大安全事件。",
        },
        {
            "block_id": "monthly-actions",
            "type": "paragraph",
            "text": "已关闭两项低风险整改。",
        },
    ],
}


@pytest.fixture
def professional_catalog(generation_db):
    from app.professional_delivery.models import (
        SkillDefinition,
        SkillVersion,
        TemplateDefinition,
        TemplateVersion,
    )

    skill = SkillDefinition(
        skill_key="security_ops_monthly_report",
        name="安全运营月报",
        category="security_operations",
        description="形成安全运营月度专业成果",
        scope_policy="both",
        status="published",
        created_by="system",
    )
    template = TemplateDefinition(
        template_key="security_ops_monthly_report",
        name="安全运营月报模板",
        purpose="安全运营月度汇报",
        deliverable_types_json=["security_ops_monthly_report"],
        scope_type="system",
        status="published",
        created_by="system",
    )
    generation_db.add_all([skill, template])
    generation_db.flush()

    template_version = TemplateVersion(
        template_id=template.id,
        version=1,
        content_hash="t" * 64,
        input_schema_json={"type": "object"},
        structure_dsl_json={"type": "document", "children": []},
        status="published",
        published_by="system",
        published_at=datetime(2026, 7, 14, 9, 0, 0),
        created_by="system",
    )
    generation_db.add(template_version)
    generation_db.flush()

    skill_version = SkillVersion(
        skill_id=skill.id,
        version=1,
        content_hash="s" * 64,
        input_schema_json={"type": "object"},
        output_schema_json={"type": "object"},
        plan_definition_json={"steps": []},
        default_template_version_id=template_version.id,
        status="published",
        published_by="system",
        published_at=datetime(2026, 7, 14, 9, 0, 0),
        created_by="system",
    )
    generation_db.add(skill_version)
    generation_db.flush()
    skill.current_published_version_id = skill_version.id
    template.current_published_version_id = template_version.id
    generation_db.commit()
    return SimpleNamespace(
        skill=skill,
        skill_version=skill_version,
        template=template,
        template_version=template_version,
    )


def _create_body(catalog, **overrides) -> dict:
    body = {
        "title": "2026 年 7 月安全运营月报",
        "deliverable_type": "security_ops_monthly_report",
        "scope_type": "personal",
        "formality": "working",
        "skill_version_uuid": catalog.skill_version.uuid,
        "template_version_uuid": catalog.template_version.uuid,
        "content": CONTENT,
        "content_summary": "安全运营月度概览",
        "creation_reason": "manual",
    }
    body.update(overrides)
    return body


def _create(client, catalog, *, key="create-monthly-report", **overrides):
    return client.post(
        "/api/ai/deliverables",
        headers={"Idempotency-Key": key},
        json=_create_body(catalog, **overrides),
    )


def _create_version(
    client,
    deliverable_uuid: str,
    *,
    key: str = "create-monthly-report-v2",
    **overrides,
):
    body = {
        "row_version": 1,
        "content": UPDATED_CONTENT,
        "content_summary": "补充巡检与整改结论",
        "change_summary": "补充巡检和整改情况",
        "creation_reason": "manual_edit",
    }
    body.update(overrides)
    return client.post(
        f"/api/ai/deliverables/{deliverable_uuid}/versions",
        headers={"Idempotency-Key": key},
        json=body,
    )


def _docx_import_bytes() -> bytes:
    document = Document()
    document.add_heading("导入后的月报", level=1)
    document.add_paragraph("导入正文")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _docx_import_bytes_with_image(*, paragraph_text: str = "图表如下") -> tuple[bytes, bytes]:
    image = (
        b"\x89PNG\r\n\x1a\n"
        b"\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
        b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89"
        b"\x00\x00\x00\x0dIDAT\x08\xd7c\xfc\xcf\xc0\xf0\x1f\x00\x05\x00\x01\xff\x89\x99=\x1d"
        b"\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    document = Document()
    document.add_heading("导入后的月报", level=1)
    document.add_paragraph(paragraph_text)
    document.add_picture(BytesIO(image))
    output = BytesIO()
    document.save(output)
    return output.getvalue(), image


def test_create_personal_deliverable_pins_versions_encrypts_content_and_replays(
    client_for_user,
    generation_db,
    professional_catalog,
) -> None:
    from app.crypto import ContentCipher, EncryptedPayload
    from app.governance_models import AuditLog
    from app.models import WorkArtifact, WorkArtifactVersion
    from app.professional_delivery.models import DeliverableIdempotencyRecord
    from app.config import get_settings

    client = client_for_user("u-1")
    first = _create(client, professional_catalog)

    assert first.status_code == 201, first.text
    payload = first.json()
    assert payload["request_id"]
    assert payload["scope_type"] == "personal"
    assert payload["project_uuid"] is None
    assert payload["owner_user_id"] == "u-1"
    assert payload["lifecycle_status"] == "draft"
    assert payload["row_version"] == 1
    assert payload["allowed_actions"] == [
        "edit",
        "update_metadata",
        "create_version",
        "manage_facts",
        "review",
        "resolve_review_issue",
        "comment",
        "reply_comment",
        "export",
    ]
    assert payload["current_version"]["version_no"] == 1
    assert payload["current_version"]["skill_version_uuid"] == professional_catalog.skill_version.uuid
    assert payload["current_version"]["template_version_uuid"] == professional_catalog.template_version.uuid
    assert payload["current_version"]["content"] == CONTENT
    assert len(payload["current_version"]["content_hash"]) == 64

    replay = _create(client, professional_catalog)
    assert replay.status_code == 201, replay.text
    assert replay.json()["deliverable_uuid"] == payload["deliverable_uuid"]

    artifact = generation_db.scalar(
        select(WorkArtifact).where(WorkArtifact.uuid == payload["deliverable_uuid"])
    )
    assert artifact is not None
    version = generation_db.scalar(
        select(WorkArtifactVersion).where(WorkArtifactVersion.artifact_id == artifact.id)
    )
    assert version is not None
    assert version.content_ciphertext
    assert CONTENT["blocks"][0]["text"].encode() not in version.content_ciphertext
    decrypted = ContentCipher(get_settings().content_encryption_key).decrypt_json(
        EncryptedPayload(version.content_ciphertext, version.content_nonce),
        version.uuid.encode("utf-8"),
    )
    assert decrypted == CONTENT
    assert generation_db.scalar(select(func.count(WorkArtifact.id))) == 1
    assert generation_db.scalar(select(func.count(WorkArtifactVersion.id))) == 1
    assert generation_db.scalar(select(func.count(DeliverableIdempotencyRecord.id))) == 1
    audits = list(
        generation_db.scalars(
            select(AuditLog).where(AuditLog.action == "professional_deliverable.create")
        )
    )
    assert len(audits) == 1
    assert audits[0].metadata_json == {
        "event": "deliverable_created",
        "status": "draft",
    }


def test_import_docx_returns_structured_content_and_audits_without_creating_version(
    client_for_user,
    generation_db,
    professional_catalog,
) -> None:
    from app.governance_models import AuditLog
    from app.models import WorkArtifactVersion

    client = client_for_user("u-1")
    created = _create(client, professional_catalog, key="docx-import-create").json()
    response = client.post(
        f"/api/ai/deliverables/{created['deliverable_uuid']}/editor/import-docx",
        headers={"Idempotency-Key": "docx-import-1"},
        files={
            "file": (
                "monthly-report.docx",
                _docx_import_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["deliverable_uuid"] == created["deliverable_uuid"]
    assert payload["source_file_name"] == "monthly-report.docx"
    assert [block["type"] for block in payload["content"]["blocks"]] == ["heading", "paragraph"]
    assert payload["warnings"] == []
    assert payload["media_count"] == 0
    assert generation_db.query(WorkArtifactVersion).count() == 1
    audit = generation_db.scalar(
        select(AuditLog).where(
            AuditLog.action == "professional_deliverable.editor.import_docx"
        )
    )
    assert audit is not None
    assert audit.metadata_json["record_count"] == 2


def test_import_docx_rejects_non_docx_filename(client_for_user, professional_catalog) -> None:
    client = client_for_user("u-1")
    created = _create(client, professional_catalog, key="docx-import-invalid-create").json()

    response = client.post(
        f"/api/ai/deliverables/{created['deliverable_uuid']}/editor/import-docx",
        files={"file": ("report.txt", b"plain text", "text/plain")},
    )

    assert response.status_code == 422
    assert response.json()["detail"]["code"] == "INVALID_DOCX_IMPORT"


def test_import_docx_persists_encrypted_media_blocks_and_replays_idempotently(
    client_for_user,
    generation_db,
    professional_catalog,
) -> None:
    from app.config import get_settings
    from app.crypto import ContentCipher, EncryptedPayload
    from app.professional_delivery.models import DeliverableMediaAsset

    client = client_for_user("docx-media-owner")
    created = _create(client, professional_catalog, key="docx-media-create").json()
    data, image = _docx_import_bytes_with_image()
    path = f"/api/ai/deliverables/{created['deliverable_uuid']}/editor/import-docx"
    first = client.post(
        path,
        headers={"Idempotency-Key": "docx-media-import-1"},
        files={"file": ("monthly-report.docx", data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    assert first.status_code == 200, first.text
    payload = first.json()
    blocks = payload["content"]["blocks"]
    assert [block["type"] for block in blocks] == ["heading", "paragraph", "image"]
    image_block = blocks[-1]
    assert image_block["asset_id"]
    assert image_block["mime_type"] == "image/png"
    assert image_block["size_bytes"] == len(image)
    asset = generation_db.query(DeliverableMediaAsset).one()
    assert asset.uuid == image_block["asset_id"]
    assert asset.status == "active"
    assert asset.content_ciphertext != image
    decrypted = ContentCipher(get_settings().content_encryption_key).decrypt_bytes(
        EncryptedPayload(asset.content_ciphertext, asset.content_nonce),
        asset.uuid.encode("utf-8"),
    )
    assert decrypted == image

    replay = client.post(
        path,
        headers={"Idempotency-Key": "docx-media-import-1"},
        files={"file": ("monthly-report.docx", data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert replay.status_code == 200, replay.text
    assert replay.json()["content"]["blocks"][-1]["asset_id"] == image_block["asset_id"]
    assert generation_db.query(DeliverableMediaAsset).count() == 1

    changed_data, _ = _docx_import_bytes_with_image(paragraph_text="图表已更新")
    conflict = client.post(
        path,
        headers={"Idempotency-Key": "docx-media-import-1"},
        files={"file": ("monthly-report.docx", changed_data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert conflict.status_code == 409, conflict.text
    assert conflict.json()["detail"]["code"] == "IDEMPOTENCY_KEY_REUSED"


def test_media_asset_upload_is_encrypted_idempotent_and_scoped(
    client_for_user,
    generation_db,
    professional_catalog,
) -> None:
    from app.governance_models import AuditLog
    from app.professional_delivery.models import DeliverableMediaAsset

    owner = client_for_user("media-owner")
    other = client_for_user("media-other")
    created = _create(owner, professional_catalog, key="media-create").json()
    deliverable_uuid = created["deliverable_uuid"]
    png = b"\x89PNG\r\n\x1a\n" + b"editor-image"
    request = {
        "files": {"file": ("cover.png", png, "image/png")},
        "headers": {"Idempotency-Key": "media-upload-1"},
    }

    uploaded = owner.post(
        f"/api/ai/deliverables/{deliverable_uuid}/editor/media",
        **request,
    )

    assert uploaded.status_code == 201, uploaded.text
    payload = uploaded.json()
    assert payload["deliverable_uuid"] == deliverable_uuid
    assert payload["media_type"] == "image/png"
    assert payload["size_bytes"] == len(png)
    assert payload["asset_uuid"]
    assert payload["download_url"].endswith(payload["asset_uuid"])
    assert payload["replayed"] is False

    replay = owner.post(
        f"/api/ai/deliverables/{deliverable_uuid}/editor/media",
        **request,
    )
    assert replay.status_code == 201, replay.text
    assert replay.json()["asset_uuid"] == payload["asset_uuid"]
    assert replay.json()["replayed"] is True
    assert generation_db.query(DeliverableMediaAsset).count() == 1
    asset = generation_db.query(DeliverableMediaAsset).one()
    assert asset.content_ciphertext != png
    assert asset.size_bytes == len(png)

    downloaded = owner.get(
        f"/api/ai/deliverables/{deliverable_uuid}/editor/media/{payload['asset_uuid']}"
    )
    assert downloaded.status_code == 200
    assert downloaded.content == png
    assert downloaded.headers["content-type"].startswith("image/png")

    preview = owner.get(
        f"/api/ai/deliverables/{deliverable_uuid}/editor/media/{payload['asset_uuid']}/preview"
    )
    assert preview.status_code == 200
    assert preview.content == png
    assert preview.headers["content-type"].startswith("image/png")

    forbidden = other.get(
        f"/api/ai/deliverables/{deliverable_uuid}/editor/media/{payload['asset_uuid']}"
    )
    assert forbidden.status_code == 404
    forbidden_preview = other.get(
        f"/api/ai/deliverables/{deliverable_uuid}/editor/media/{payload['asset_uuid']}/preview"
    )
    assert forbidden_preview.status_code == 404

    deleted = owner.delete(
        f"/api/ai/deliverables/{deliverable_uuid}/editor/media/{payload['asset_uuid']}"
    )
    assert deleted.status_code == 200, deleted.text
    assert deleted.json()["status"] == "deleted"
    assert owner.get(
        f"/api/ai/deliverables/{deliverable_uuid}/editor/media/{payload['asset_uuid']}"
    ).status_code == 404
    assert owner.get(
        f"/api/ai/deliverables/{deliverable_uuid}/editor/media/{payload['asset_uuid']}/preview"
    ).status_code == 404
    assert owner.delete(
        f"/api/ai/deliverables/{deliverable_uuid}/editor/media/{payload['asset_uuid']}"
    ).status_code == 404
    audits = list(
        generation_db.scalars(
            select(AuditLog).where(
                AuditLog.action.in_(
                    [
                        "professional_deliverable.editor.media.upload",
                        "professional_deliverable.editor.media.preview",
                        "professional_deliverable.editor.media.delete",
                    ]
                )
            )
        )
    )
    assert len(audits) == 3
    upload_audit = next(
        item for item in audits if item.action == "professional_deliverable.editor.media.upload"
    )
    assert upload_audit.metadata_json["size_bytes"] == len(png)


def test_media_asset_rejects_mismatched_signature_and_oversized_payload(
    client_for_user,
    professional_catalog,
    monkeypatch,
) -> None:
    import app.professional_delivery.media_service as media_service

    owner = client_for_user("media-validation-owner")
    created = _create(owner, professional_catalog, key="media-validation-create").json()
    url = f"/api/ai/deliverables/{created['deliverable_uuid']}/editor/media"

    mismatch = owner.post(
        url,
        headers={"Idempotency-Key": "media-mismatch"},
        files={"file": ("not-a-png.png", b"plain text", "image/png")},
    )
    assert mismatch.status_code == 422
    assert mismatch.json()["detail"]["code"] == "INVALID_MEDIA_ASSET"

    monkeypatch.setattr(media_service, "MAX_MEDIA_ASSET_BYTES", 4)
    oversized = owner.post(
        url,
        headers={"Idempotency-Key": "media-oversized"},
        files={"file": ("large.png", b"\x89PNG\r\n\x1a\n", "image/png")},
    )
    assert oversized.status_code == 413
    assert oversized.json()["detail"]["code"] == "MEDIA_ASSET_TOO_LARGE"


def test_media_lifecycle_counts_references_and_only_cleans_orphans(
    client_for_user,
    generation_db,
    professional_catalog,
) -> None:
    from app.config import get_settings
    from app.crypto import ContentCipher
    from app.models import WorkArtifact
    from app.professional_delivery.media_service import (
        cleanup_orphaned_media_assets,
        count_media_asset_references,
    )
    from app.professional_delivery.models import DeliverableMediaAsset

    client = client_for_user("media-lifecycle-owner")
    created = _create(client, professional_catalog, key="media-lifecycle-create").json()
    deliverable_uuid = created["deliverable_uuid"]
    image = b"\x89PNG\r\n\x1a\n" + b"lifecycle-image"
    first = client.post(
        f"/api/ai/deliverables/{deliverable_uuid}/editor/media",
        files={"file": ("referenced.png", image, "image/png")},
        headers={"Idempotency-Key": "media-lifecycle-referenced"},
    ).json()
    second = client.post(
        f"/api/ai/deliverables/{deliverable_uuid}/editor/media",
        files={"file": ("orphan.png", image, "image/png")},
        headers={"Idempotency-Key": "media-lifecycle-orphan"},
    ).json()

    version = _create_version(
        client,
        deliverable_uuid,
        key="media-lifecycle-version",
        content={
            "schema_version": "1",
            "blocks": [{
                "block_id": "hero-image",
                "type": "image",
                "asset_id": first["asset_uuid"],
                "mime_type": "image/png",
                "size_bytes": len(image),
            }],
        },
    )
    assert version.status_code == 201, version.text

    cipher = ContentCipher(get_settings().content_encryption_key)
    artifact = generation_db.scalar(
        select(WorkArtifact).where(WorkArtifact.uuid == deliverable_uuid)
    )
    assert artifact is not None
    assert count_media_asset_references(
        generation_db,
        deliverable_id=artifact.id,
        asset_uuid=first["asset_uuid"],
        cipher=cipher,
    ) == 1
    first_row = generation_db.scalar(
        select(DeliverableMediaAsset).where(DeliverableMediaAsset.uuid == first["asset_uuid"])
    )
    second_row = generation_db.scalar(
        select(DeliverableMediaAsset).where(DeliverableMediaAsset.uuid == second["asset_uuid"])
    )
    assert first_row is not None and second_row is not None
    first_row.created_at = second_row.created_at = datetime(2026, 1, 1)
    deleted = cleanup_orphaned_media_assets(
        generation_db,
        deliverable_id=artifact.id,
        cipher=cipher,
        older_than=datetime(2026, 2, 1),
    )
    assert deleted == [second["asset_uuid"]]
    assert first_row.status == "active"
    assert second_row.status == "deleted"


def test_media_scan_quarantines_tampered_ciphertext(
    client_for_user,
    generation_db,
    professional_catalog,
) -> None:
    from app.config import get_settings
    from app.crypto import ContentCipher
    from app.professional_delivery.media_service import scan_media_asset
    from app.professional_delivery.models import DeliverableMediaAsset

    client = client_for_user("media-scan-owner")
    created = _create(client, professional_catalog, key="media-scan-create").json()
    payload = client.post(
        f"/api/ai/deliverables/{created['deliverable_uuid']}/editor/media",
        files={"file": ("scan.png", b"\x89PNG\r\n\x1a\n" + b"scan", "image/png")},
        headers={"Idempotency-Key": "media-scan-upload"},
    ).json()
    asset = generation_db.scalar(
        select(DeliverableMediaAsset).where(DeliverableMediaAsset.uuid == payload["asset_uuid"])
    )
    assert asset is not None
    asset.content_ciphertext = b"tampered"
    valid = scan_media_asset(
        generation_db,
        asset=asset,
        cipher=ContentCipher(get_settings().content_encryption_key),
    )
    assert valid is False
    assert asset.status == "quarantined"


def test_editor_rejects_media_asset_from_another_deliverable(
    client_for_user,
    professional_catalog,
) -> None:
    owner = client_for_user("media-reference-owner")
    source = _create(owner, professional_catalog, key="media-reference-source").json()
    target = _create(owner, professional_catalog, key="media-reference-target").json()
    png = b"\x89PNG\r\n\x1a\n" + b"editor-image"
    uploaded = owner.post(
        f"/api/ai/deliverables/{source['deliverable_uuid']}/editor/media",
        headers={"Idempotency-Key": "media-reference-upload"},
        files={"file": ("cover.png", png, "image/png")},
    )
    assert uploaded.status_code == 201, uploaded.text
    asset = uploaded.json()

    version_uuid = target["current_version"]["version_uuid"]
    lease = owner.post(
        f"/api/ai/deliverables/{target['deliverable_uuid']}/draft/lease",
        json={"row_version": 1, "base_version_uuid": version_uuid},
    )
    assert lease.status_code == 200, lease.text
    rejected = owner.put(
        f"/api/ai/deliverables/{target['deliverable_uuid']}/draft",
        headers={"Idempotency-Key": "media-reference-save"},
        json={
            "row_version": 1,
            "base_version_uuid": version_uuid,
            "draft_revision": 0,
            "content": {
                "schema_version": "2",
                "blocks": [{
                    "block_id": "cover",
                    "type": "media",
                    "asset_id": asset["asset_uuid"],
                    "mime_type": "image/png",
                    "size_bytes": len(png),
                }],
            },
            "fencing_token": lease.json()["fencing_token"],
        },
    )
    assert rejected.status_code == 422, rejected.text
    assert rejected.json()["detail"]["code"] == "INVALID_DELIVERABLE_CONTENT"


def test_update_metadata_uses_optimistic_lock_and_preserves_version_snapshot(
    client_for_user,
    generation_db,
    professional_catalog,
) -> None:
    from app.governance_models import AuditLog
    from app.models import WorkArtifact, WorkArtifactVersion
    from app.professional_delivery.models import DeliverableIdempotencyRecord

    client = client_for_user("u-1")
    created = _create(client, professional_catalog, key="metadata-create").json()
    deliverable_uuid = created["deliverable_uuid"]
    original_snapshot = created["current_version"]["title_snapshot"]
    body = {"row_version": created["row_version"], "title": "安全运营月报（终稿）"}

    updated = client.patch(
        f"/api/ai/deliverables/{deliverable_uuid}",
        headers={"Idempotency-Key": "metadata-update"},
        json=body,
    )

    assert updated.status_code == 200, updated.text
    payload = updated.json()
    assert payload["title"] == "安全运营月报（终稿）"
    assert payload["row_version"] == 2
    assert payload["current_version"]["title_snapshot"] == original_snapshot

    replay = client.patch(
        f"/api/ai/deliverables/{deliverable_uuid}",
        headers={"Idempotency-Key": "metadata-update"},
        json=body,
    )
    assert replay.status_code == 200, replay.text
    assert replay.json()["row_version"] == 2

    stale = client.patch(
        f"/api/ai/deliverables/{deliverable_uuid}",
        headers={"Idempotency-Key": "metadata-update-stale"},
        json={"row_version": 1, "title": "过期写入"},
    )
    assert stale.status_code == 409
    assert stale.json()["detail"]["code"] == "DELIVERABLE_VERSION_CONFLICT"

    artifact = generation_db.scalar(
        select(WorkArtifact).where(WorkArtifact.uuid == deliverable_uuid)
    )
    assert artifact is not None
    version = generation_db.get(WorkArtifactVersion, artifact.current_version_id)
    assert version is not None
    assert artifact.title == "安全运营月报（终稿）"
    assert version.title_snapshot == original_snapshot
    assert generation_db.scalar(
        select(func.count(DeliverableIdempotencyRecord.id))
    ) == 2
    assert generation_db.scalar(
        select(func.count(AuditLog.id)).where(
            AuditLog.action == "professional_deliverable.metadata.update"
        )
    ) == 1


def test_delivered_result_submits_only_deidentified_exact_version_experience_candidate(
    client_for_user,
    generation_db,
    professional_catalog,
) -> None:
    from app.config import get_settings
    from app.crypto import ContentCipher, EncryptedPayload
    from app.governance_models import AuditLog
    from app.models import ExperienceLibrary, WorkArtifact, WorkArtifactVersion
    from app.professional_delivery.models import DeliverableExperienceCandidate

    client = client_for_user("u-1")
    created = _create(client, professional_catalog, key="experience-create").json()
    artifact = generation_db.scalar(
        select(WorkArtifact).where(
            WorkArtifact.uuid == created["deliverable_uuid"]
        )
    )
    assert artifact is not None
    version = generation_db.get(WorkArtifactVersion, artifact.current_version_id)
    assert version is not None
    artifact.lifecycle_status = "delivered"
    artifact.delivered_version_id = version.id
    generation_db.commit()

    body = {
        "row_version": artifact.row_version,
        "version_uuid": version.uuid,
        "content_hash": version.content_hash,
        "candidate_type": "rule",
        "deidentified_summary": "正式成果提交前，应先完成事实证据和敏感信息复核。",
    }
    submitted = client.post(
        f"/api/ai/deliverables/{artifact.uuid}/experience-candidates",
        headers={"Idempotency-Key": "experience-submit"},
        json=body,
    )

    assert submitted.status_code == 201, submitted.text
    payload = submitted.json()["candidate"]
    assert payload["candidate_type"] == "rule"
    assert payload["status"] == "pending_review"
    assert payload["source_scope_type"] == "personal"
    assert payload["source_project_uuid"] is None
    assert payload["version_uuid"] == version.uuid
    assert payload["content_hash"] == version.content_hash
    assert payload["deidentified_summary"] == body["deidentified_summary"]

    replay = client.post(
        f"/api/ai/deliverables/{artifact.uuid}/experience-candidates",
        headers={"Idempotency-Key": "experience-submit"},
        json=body,
    )
    assert replay.status_code == 201, replay.text
    assert replay.json()["candidate"]["candidate_uuid"] == payload["candidate_uuid"]

    stale = client.post(
        f"/api/ai/deliverables/{artifact.uuid}/experience-candidates",
        headers={"Idempotency-Key": "experience-stale"},
        json={**body, "content_hash": "0" * 64},
    )
    assert stale.status_code == 409
    assert stale.json()["detail"]["code"] == "EXPERIENCE_CANDIDATE_TARGET_STALE"

    sensitive = client.post(
        f"/api/ai/deliverables/{artifact.uuid}/experience-candidates",
        headers={"Idempotency-Key": "experience-sensitive"},
        json={**body, "deidentified_summary": "联系 user@example.com 复核规则。"},
    )
    assert sensitive.status_code == 422
    assert (
        sensitive.json()["detail"]["code"]
        == "EXPERIENCE_CANDIDATE_NOT_DEIDENTIFIED"
    )

    candidate = generation_db.scalar(select(DeliverableExperienceCandidate))
    assert candidate is not None
    assert body["deidentified_summary"].encode() not in candidate.payload_ciphertext
    decrypted = ContentCipher(get_settings().content_encryption_key).decrypt_json(
        EncryptedPayload(candidate.payload_ciphertext, candidate.payload_nonce),
        candidate.uuid.encode("utf-8"),
    )
    assert decrypted == {"deidentified_summary": body["deidentified_summary"]}
    assert generation_db.scalar(select(func.count(DeliverableExperienceCandidate.id))) == 1
    assert generation_db.scalar(select(func.count(ExperienceLibrary.id))) == 0
    assert generation_db.scalar(
        select(func.count(AuditLog.id)).where(
            AuditLog.action
            == "professional_deliverable.experience_candidate.create"
        )
    ) == 1


def test_personal_deliverable_is_hidden_from_other_users(
    client_for_user,
    professional_catalog,
) -> None:
    owner = client_for_user("u-1")
    outsider = client_for_user("u-2")
    created = _create(owner, professional_catalog).json()

    owner_list = owner.get("/api/ai/deliverables")
    assert owner_list.status_code == 200
    assert [item["deliverable_uuid"] for item in owner_list.json()["items"]] == [
        created["deliverable_uuid"]
    ]
    outsider_list = outsider.get("/api/ai/deliverables")
    assert outsider_list.status_code == 200
    assert outsider_list.json()["items"] == []
    assert outsider.get(
        f"/api/ai/deliverables/{created['deliverable_uuid']}"
    ).status_code == 404


def test_project_deliverable_uses_membership_instead_of_creator_ownership(
    client_for_user,
    professional_catalog,
) -> None:
    owner = client_for_user("u-1")
    member = client_for_user("u-2")
    outsider = client_for_user("u-3")
    project = owner.post(
        "/api/ai/projects",
        json={"name": "安全运营项目", "description": "月度服务"},
    ).json()
    assert owner.post(
        f"/api/ai/projects/{project['project_uuid']}/members",
        json={"user_id": "u-2", "role": "member"},
    ).status_code == 201

    response = _create(
        owner,
        professional_catalog,
        key="create-project-monthly-report",
        scope_type="project",
        formality="formal",
        project_uuid=project["project_uuid"],
    )
    assert response.status_code == 201, response.text
    created = response.json()
    assert created["project_uuid"] == project["project_uuid"]
    member_detail = member.get(
        f"/api/ai/deliverables/{created['deliverable_uuid']}"
    )
    assert member_detail.status_code == 200
    assert member_detail.json()["allowed_actions"] == [
        "edit",
        "update_metadata",
        "create_version",
        "manage_facts",
        "reply_comment",
        "export",
    ]
    assert [
        item["deliverable_uuid"]
        for item in member.get("/api/ai/deliverables").json()["items"]
    ] == [created["deliverable_uuid"]]
    assert outsider.get(
        f"/api/ai/deliverables/{created['deliverable_uuid']}"
    ).status_code == 404
    assert outsider.get("/api/ai/deliverables").json()["items"] == []


def test_read_only_project_member_cannot_create_deliverable(
    client_for_user,
    professional_catalog,
) -> None:
    owner = client_for_user("u-1")
    read_only = client_for_user("u-2")
    project = owner.post(
        "/api/ai/projects",
        json={"name": "只读项目", "description": ""},
    ).json()
    owner.post(
        f"/api/ai/projects/{project['project_uuid']}/members",
        json={"user_id": "u-2", "role": "read_only"},
    )

    response = _create(
        read_only,
        professional_catalog,
        key="read-only-create",
        scope_type="project",
        project_uuid=project["project_uuid"],
    )

    assert response.status_code == 403
    assert response.json()["detail"]["code"] == "PROJECT_DELIVERABLE_WRITE_FORBIDDEN"


def test_create_rejects_missing_idempotency_key_and_invalid_catalog_or_content(
    client_for_user,
    generation_db,
    professional_catalog,
) -> None:
    from app.models import WorkArtifact

    client = client_for_user("u-1")
    missing_key = client.post(
        "/api/ai/deliverables",
        json=_create_body(professional_catalog),
    )
    assert missing_key.status_code == 400
    assert missing_key.json()["detail"]["code"] == "IDEMPOTENCY_KEY_REQUIRED"

    professional_catalog.skill_version.status = "retired"
    generation_db.commit()
    unavailable = _create(client, professional_catalog, key="unavailable-skill")
    assert unavailable.status_code == 422
    assert unavailable.json()["detail"]["code"] == "SKILL_VERSION_NOT_AVAILABLE"
    professional_catalog.skill_version.status = "published"
    generation_db.commit()

    invalid_content = {
        "schema_version": "1",
        "blocks": [
            {"block_id": "duplicate", "type": "paragraph", "text": "一"},
            {"block_id": "duplicate", "type": "paragraph", "text": "二"},
        ],
    }
    invalid = _create(
        client,
        professional_catalog,
        key="invalid-content",
        content=invalid_content,
    )
    assert invalid.status_code == 422
    assert invalid.json()["detail"]["code"] == "INVALID_DELIVERABLE_CONTENT"
    assert generation_db.scalar(select(func.count(WorkArtifact.id))) == 0


@pytest.mark.parametrize(
    "media_block",
    [
        {"block_id": "hero", "type": "image", "url": "javascript:alert(1)"},
        {"block_id": "hero", "type": "image", "url": "data:image/png;base64,AAAA"},
        {"block_id": "hero", "type": "image", "url": "https://example.com/hero.png"},
        {"block_id": "hero", "type": "image", "asset_id": "asset-1", "mime_type": "image/svg+xml"},
        {"block_id": "hero", "type": "image", "asset_id": "asset-1", "mime_type": "image/png", "size_bytes": 11 * 1024 * 1024},
    ],
)
def test_create_rejects_unsafe_or_unbounded_media_blocks(
    client_for_user,
    professional_catalog,
    media_block,
) -> None:
    client = client_for_user("u-1")
    response = _create(
        client,
        professional_catalog,
        key=f"unsafe-media-{media_block.get('mime_type', 'url')}-{media_block.get('size_bytes', 0)}",
        content={"schema_version": "2", "blocks": [media_block]},
    )
    assert response.status_code == 422
    assert response.json()["detail"]["code"] == "INVALID_DELIVERABLE_CONTENT"


def test_create_accepts_internal_media_asset_reference(
    client_for_user,
    professional_catalog,
) -> None:
    client = client_for_user("u-1")
    response = _create(
        client,
        professional_catalog,
        key="safe-media-asset",
        content={
            "schema_version": "2",
            "blocks": [
                {
                    "block_id": "hero",
                    "type": "image",
                    "asset_id": "asset-1",
                    "mime_type": "image/png",
                    "size_bytes": 1024,
                    "alt": "系统架构图",
                }
            ],
        },
    )
    assert response.status_code == 201, response.text


def test_skill_scope_policy_is_enforced(
    client_for_user,
    generation_db,
    professional_catalog,
) -> None:
    owner = client_for_user("u-1")
    project = owner.post(
        "/api/ai/projects",
        json={"name": "项目", "description": ""},
    ).json()
    professional_catalog.skill.scope_policy = "personal"
    generation_db.commit()

    response = _create(
        owner,
        professional_catalog,
        key="wrong-scope",
        scope_type="project",
        project_uuid=project["project_uuid"],
    )

    assert response.status_code == 422
    assert response.json()["detail"]["code"] == "SKILL_SCOPE_MISMATCH"


def test_audit_failure_rolls_back_deliverable_version_and_idempotency(
    client_for_user,
    generation_db,
    professional_catalog,
    monkeypatch,
) -> None:
    from app.models import WorkArtifact, WorkArtifactVersion
    from app.professional_delivery import routes
    from app.professional_delivery.models import DeliverableIdempotencyRecord

    def fail_audit(*_args, **_kwargs) -> None:
        raise RuntimeError("audit unavailable")

    monkeypatch.setattr(routes, "write_request_audit", fail_audit)
    client = client_for_user("u-1")
    with pytest.raises(RuntimeError, match="audit unavailable"):
        _create(client, professional_catalog, key="audit-failure")

    assert generation_db.scalar(select(func.count(WorkArtifact.id))) == 0
    assert generation_db.scalar(select(func.count(WorkArtifactVersion.id))) == 0
    assert generation_db.scalar(select(func.count(DeliverableIdempotencyRecord.id))) == 0


def test_create_deliverable_version_is_immutable_encrypted_audited_and_idempotent(
    client_for_user,
    generation_db,
    professional_catalog,
) -> None:
    from app.config import get_settings
    from app.crypto import ContentCipher, EncryptedPayload
    from app.governance_models import AuditLog
    from app.models import WorkArtifact, WorkArtifactVersion
    from app.professional_delivery.models import DeliverableIdempotencyRecord

    client = client_for_user("u-1")
    created = _create(client, professional_catalog).json()
    artifact = generation_db.scalar(
        select(WorkArtifact).where(
            WorkArtifact.uuid == created["deliverable_uuid"]
        )
    )
    assert artifact is not None
    original = generation_db.get(WorkArtifactVersion, artifact.current_version_id)
    assert original is not None
    original_state = (
        original.content_ciphertext,
        original.content_nonce,
        original.content_hash,
        original.title_snapshot,
        original.summary_snapshot,
    )

    response = _create_version(client, created["deliverable_uuid"])

    assert response.status_code == 201, response.text
    payload = response.json()
    assert payload["deliverable_uuid"] == created["deliverable_uuid"]
    assert payload["version"]["version_no"] == 2
    assert payload["version"]["parent_version_uuid"] == original.uuid
    assert payload["version"]["skill_version_uuid"] == professional_catalog.skill_version.uuid
    assert payload["version"]["template_version_uuid"] == professional_catalog.template_version.uuid
    assert payload["version"]["content"] == UPDATED_CONTENT
    assert payload["version"]["change_summary"] == "补充巡检和整改情况"
    assert payload["version"]["creation_reason"] == "manual_edit"
    assert payload["version"]["content_hash"] != original.content_hash

    replay = _create_version(client, created["deliverable_uuid"])
    assert replay.status_code == 201, replay.text
    assert replay.json()["version"] == payload["version"]

    generation_db.refresh(artifact)
    versions = list(
        generation_db.scalars(
            select(WorkArtifactVersion)
            .where(WorkArtifactVersion.artifact_id == artifact.id)
            .order_by(WorkArtifactVersion.version)
        )
    )
    assert len(versions) == 2
    assert (
        versions[0].content_ciphertext,
        versions[0].content_nonce,
        versions[0].content_hash,
        versions[0].title_snapshot,
        versions[0].summary_snapshot,
    ) == original_state
    assert versions[1].parent_version_id == versions[0].id
    assert versions[1].content_ciphertext
    assert UPDATED_CONTENT["blocks"][0]["text"].encode() not in versions[1].content_ciphertext
    decrypted = ContentCipher(get_settings().content_encryption_key).decrypt_json(
        EncryptedPayload(versions[1].content_ciphertext, versions[1].content_nonce),
        versions[1].uuid.encode("utf-8"),
    )
    assert decrypted == UPDATED_CONTENT
    assert artifact.version == 2
    assert artifact.current_version_id == versions[1].id
    assert artifact.row_version == 2
    assert artifact.lifecycle_status == "draft"
    assert generation_db.scalar(select(func.count(DeliverableIdempotencyRecord.id))) == 2
    audits = list(
        generation_db.scalars(
            select(AuditLog).where(
                AuditLog.action == "professional_deliverable.version.create"
            )
        )
    )
    assert len(audits) == 1
    assert audits[0].metadata_json == {
        "event": "deliverable_version_created",
        "from_version": 1,
        "to_version": 2,
        "status": "draft",
    }


def test_structured_draft_lease_save_and_commit_preserve_immutable_version_boundary(
    client_for_user,
    generation_db,
    professional_catalog,
) -> None:
    from app.models import WorkArtifact, WorkArtifactVersion

    client = client_for_user("u-1")
    created = _create(client, professional_catalog, key="editor-create").json()
    deliverable_uuid = created["deliverable_uuid"]
    version_uuid = created["current_version"]["version_uuid"]

    draft_response = client.get(f"/api/ai/deliverables/{deliverable_uuid}/draft")
    assert draft_response.status_code == 200, draft_response.text
    draft = draft_response.json()
    assert draft["draft_revision"] == 0
    assert draft["base_version_uuid"] == version_uuid

    lease_response = client.post(
        f"/api/ai/deliverables/{deliverable_uuid}/draft/lease",
        json={"row_version": 1, "base_version_uuid": version_uuid},
    )
    assert lease_response.status_code == 200, lease_response.text
    fencing_token = lease_response.json()["fencing_token"]

    edited_content = {
        "schema_version": "2",
        "blocks": [
            {"block_id": "monthly-overview", "type": "paragraph", "text": "已完成巡检。"},
            {"block_id": "monthly-actions", "type": "paragraph", "text": "已关闭整改项。"},
        ],
    }
    saved = client.put(
        f"/api/ai/deliverables/{deliverable_uuid}/draft",
        headers={"Idempotency-Key": "editor-draft-save"},
        json={
            "row_version": 1,
            "base_version_uuid": version_uuid,
            "draft_revision": 0,
            "content": edited_content,
            "content_summary": "已完成巡检并关闭整改项",
            "fencing_token": fencing_token,
        },
    )
    assert saved.status_code == 200, saved.text
    assert saved.json()["draft_revision"] == 1
    assert saved.json()["content"] == edited_content

    replayed_save = client.put(
        f"/api/ai/deliverables/{deliverable_uuid}/draft",
        headers={"Idempotency-Key": "editor-draft-save"},
        json={
            "row_version": 1,
            "base_version_uuid": version_uuid,
            "draft_revision": 0,
            "content": edited_content,
            "content_summary": "已完成巡检并关闭整改项",
            "fencing_token": fencing_token,
        },
    )
    assert replayed_save.status_code == 200, replayed_save.text
    assert replayed_save.json()["draft_revision"] == 1

    reused_key = client.put(
        f"/api/ai/deliverables/{deliverable_uuid}/draft",
        headers={"Idempotency-Key": "editor-draft-save"},
        json={
            "row_version": 1,
            "base_version_uuid": version_uuid,
            "draft_revision": 0,
            "content": {**edited_content, "blocks": edited_content["blocks"][:1]},
            "fencing_token": fencing_token,
        },
    )
    assert reused_key.status_code == 409, reused_key.text
    assert reused_key.json()["detail"]["code"] == "IDEMPOTENCY_KEY_REUSED"

    committed = client.post(
        f"/api/ai/deliverables/{deliverable_uuid}/draft/commit",
        headers={"Idempotency-Key": "editor-draft-commit"},
        json={
            "row_version": 1,
            "base_version_uuid": version_uuid,
            "draft_revision": 1,
            "change_summary": "结构化编辑器保存",
            "creation_reason": "manual_edit",
            "fencing_token": fencing_token,
        },
    )
    assert committed.status_code == 201, committed.text
    committed_payload = committed.json()
    assert committed_payload["version"]["version_no"] == 2
    assert committed_payload["version"]["content"] == edited_content

    artifact = generation_db.scalar(
        select(WorkArtifact).where(WorkArtifact.uuid == deliverable_uuid)
    )
    assert artifact is not None
    versions = list(
        generation_db.scalars(
            select(WorkArtifactVersion)
            .where(WorkArtifactVersion.artifact_id == artifact.id)
            .order_by(WorkArtifactVersion.version)
        )
    )
    assert len(versions) == 2
    assert versions[0].content_hash != versions[1].content_hash
    assert artifact.row_version == 2

    stale = client.put(
        f"/api/ai/deliverables/{deliverable_uuid}/draft",
        headers={"Idempotency-Key": "editor-draft-stale"},
        json={
            "row_version": 1,
            "base_version_uuid": version_uuid,
            "draft_revision": 0,
            "content": edited_content,
        },
    )
    assert stale.status_code == 409
    assert stale.json()["detail"]["code"] == "DELIVERABLE_DRAFT_CONFLICT"


def test_version_idempotency_replay_returns_original_version_after_later_write(
    client_for_user,
    generation_db,
    professional_catalog,
) -> None:
    from app.models import WorkArtifactVersion

    client = client_for_user("u-1")
    created = _create(client, professional_catalog).json()
    first = _create_version(
        client,
        created["deliverable_uuid"],
        key="original-v2-write",
    )
    assert first.status_code == 201, first.text
    first_version = first.json()["version"]

    later_content = {
        "schema_version": "1",
        "blocks": [
            {
                "block_id": "overview",
                "type": "paragraph",
                "text": "第三版月报正文",
            }
        ],
    }
    later = _create_version(
        client,
        created["deliverable_uuid"],
        key="later-v3-write",
        row_version=2,
        content=later_content,
        change_summary="形成第三版",
    )
    assert later.status_code == 201, later.text
    assert later.json()["version"]["version_no"] == 3

    replay = _create_version(
        client,
        created["deliverable_uuid"],
        key="original-v2-write",
    )

    assert replay.status_code == 201, replay.text
    assert replay.json()["version"] == first_version
    assert generation_db.scalar(select(func.count(WorkArtifactVersion.id))) == 3


def test_create_deliverable_version_rejects_stale_row_and_idempotency_reuse(
    client_for_user,
    generation_db,
    professional_catalog,
) -> None:
    from app.models import WorkArtifactVersion

    client = client_for_user("u-1")
    created = _create(client, professional_catalog).json()
    first = _create_version(client, created["deliverable_uuid"])
    assert first.status_code == 201, first.text

    stale = _create_version(
        client,
        created["deliverable_uuid"],
        key="stale-version-write",
        row_version=1,
        change_summary="过期客户端写入",
    )
    assert stale.status_code == 409
    assert stale.json()["detail"] == {
        "code": "DELIVERABLE_VERSION_CONFLICT",
        "message": "成果已被其他操作更新，请刷新后重试",
        "current_row_version": 2,
        "current_version_no": 2,
    }

    reused = _create_version(
        client,
        created["deliverable_uuid"],
        key="create-monthly-report-v2",
        row_version=2,
        change_summary="用相同幂等键提交不同请求",
    )
    assert reused.status_code == 409
    assert reused.json()["detail"]["code"] == "IDEMPOTENCY_KEY_REUSED"
    assert generation_db.scalar(select(func.count(WorkArtifactVersion.id))) == 2


def test_project_deliverable_version_requires_active_writer_membership(
    client_for_user,
    professional_catalog,
) -> None:
    owner = client_for_user("u-1")
    member = client_for_user("u-2")
    read_only = client_for_user("u-3")
    outsider = client_for_user("u-4")
    project = owner.post(
        "/api/ai/projects",
        json={"name": "版本权限项目", "description": ""},
    ).json()
    owner.post(
        f"/api/ai/projects/{project['project_uuid']}/members",
        json={"user_id": "u-2", "role": "member"},
    )
    owner.post(
        f"/api/ai/projects/{project['project_uuid']}/members",
        json={"user_id": "u-3", "role": "read_only"},
    )
    created = _create(
        owner,
        professional_catalog,
        key="create-project-version-report",
        scope_type="project",
        formality="formal",
        project_uuid=project["project_uuid"],
    ).json()

    forbidden = _create_version(
        read_only,
        created["deliverable_uuid"],
        key="read-only-version-write",
    )
    assert forbidden.status_code == 403
    assert forbidden.json()["detail"]["code"] == "PROJECT_DELIVERABLE_WRITE_FORBIDDEN"
    assert _create_version(
        outsider,
        created["deliverable_uuid"],
        key="outsider-version-write",
    ).status_code == 404

    allowed = _create_version(
        member,
        created["deliverable_uuid"],
        key="member-version-write",
    )
    assert allowed.status_code == 201, allowed.text
    assert allowed.json()["version"]["version_no"] == 2


def test_create_revision_preserves_delivery_milestones_and_audit_failure_rolls_back(
    client_for_user,
    generation_db,
    professional_catalog,
    monkeypatch,
) -> None:
    from app.models import WorkArtifact, WorkArtifactVersion
    from app.professional_delivery import routes
    from app.professional_delivery.models import DeliverableIdempotencyRecord

    client = client_for_user("u-1")
    created = _create(client, professional_catalog).json()
    artifact = generation_db.scalar(
        select(WorkArtifact).where(
            WorkArtifact.uuid == created["deliverable_uuid"]
        )
    )
    assert artifact is not None
    delivered_version_id = artifact.current_version_id
    artifact.lifecycle_status = "delivered"
    artifact.approved_version_id = delivered_version_id
    artifact.delivered_version_id = delivered_version_id
    generation_db.commit()

    def fail_audit(*_args, **_kwargs) -> None:
        raise RuntimeError("audit unavailable")

    monkeypatch.setattr(routes, "write_request_audit", fail_audit)
    with pytest.raises(RuntimeError, match="audit unavailable"):
        _create_version(
            client,
            created["deliverable_uuid"],
            key="revision-audit-failure",
        )

    generation_db.refresh(artifact)
    assert artifact.lifecycle_status == "delivered"
    assert artifact.row_version == 1
    assert artifact.current_version_id == delivered_version_id
    assert artifact.approved_version_id == delivered_version_id
    assert artifact.delivered_version_id == delivered_version_id
    assert generation_db.scalar(select(func.count(WorkArtifactVersion.id))) == 1
    assert generation_db.scalar(select(func.count(DeliverableIdempotencyRecord.id))) == 1


def test_delivered_deliverable_can_create_revision_without_overwriting_milestones(
    client_for_user,
    generation_db,
    professional_catalog,
) -> None:
    from app.models import WorkArtifact, WorkArtifactVersion

    client = client_for_user("u-1")
    created = _create(client, professional_catalog).json()
    artifact = generation_db.scalar(
        select(WorkArtifact).where(
            WorkArtifact.uuid == created["deliverable_uuid"]
        )
    )
    assert artifact is not None
    delivered_version_id = artifact.current_version_id
    artifact.lifecycle_status = "delivered"
    artifact.approved_version_id = delivered_version_id
    artifact.delivered_version_id = delivered_version_id
    generation_db.commit()

    response = _create_version(
        client,
        created["deliverable_uuid"],
        key="delivered-revision-success",
    )

    assert response.status_code == 201, response.text
    assert response.json()["version"]["version_no"] == 2
    generation_db.refresh(artifact)
    assert artifact.lifecycle_status == "draft"
    assert artifact.current_version_id != delivered_version_id
    assert artifact.approved_version_id == delivered_version_id
    assert artifact.delivered_version_id == delivered_version_id
    assert generation_db.scalar(select(func.count(WorkArtifactVersion.id))) == 2


def test_version_history_is_paginated_and_exact_version_restores_immutable_content(
    client_for_user,
    professional_catalog,
) -> None:
    owner = client_for_user("u-1")
    outsider = client_for_user("u-2")
    created = _create(owner, professional_catalog).json()
    version_1 = created["current_version"]
    version_2_response = _create_version(owner, created["deliverable_uuid"])
    assert version_2_response.status_code == 201, version_2_response.text
    version_2 = version_2_response.json()["version"]
    version_3_response = _create_version(
        owner,
        created["deliverable_uuid"],
        key="history-v3",
        row_version=2,
        content={
            "schema_version": "1",
            "blocks": [
                {
                    "block_id": "monthly-overview",
                    "type": "paragraph",
                    "text": "第三版月报正文",
                }
            ],
        },
        change_summary="形成第三版",
    )
    assert version_3_response.status_code == 201, version_3_response.text
    version_3 = version_3_response.json()["version"]

    first_page = owner.get(
        f"/api/ai/deliverables/{created['deliverable_uuid']}/versions",
        params={"page": 1, "page_size": 2},
    )
    assert first_page.status_code == 200, first_page.text
    history = first_page.json()
    assert history["request_id"]
    assert history["deliverable_uuid"] == created["deliverable_uuid"]
    assert history["total"] == 3
    assert history["page"] == 1
    assert history["page_size"] == 2
    assert [item["version_no"] for item in history["items"]] == [3, 2]
    assert history["items"][0]["is_current"] is True
    assert history["items"][1]["is_current"] is False
    assert "content" not in history["items"][0]

    second_page = owner.get(
        f"/api/ai/deliverables/{created['deliverable_uuid']}/versions",
        params={"page": 2, "page_size": 2},
    )
    assert second_page.status_code == 200, second_page.text
    assert [item["version_no"] for item in second_page.json()["items"]] == [1]

    exact = owner.get(
        f"/api/ai/deliverables/{created['deliverable_uuid']}"
        f"/versions/{version_1['version_uuid']}"
    )
    assert exact.status_code == 200, exact.text
    exact_payload = exact.json()
    assert exact_payload["request_id"]
    assert exact_payload["deliverable_uuid"] == created["deliverable_uuid"]
    assert exact_payload["version"] == version_1
    assert exact_payload["version"]["content"] == CONTENT
    assert exact_payload["version"]["content_hash"] != version_2["content_hash"]
    assert exact_payload["version"]["content_hash"] != version_3["content_hash"]

    other = _create(
        owner,
        professional_catalog,
        key="other-deliverable-for-version-injection",
        title="另一个成果",
    ).json()
    injected = owner.get(
        f"/api/ai/deliverables/{created['deliverable_uuid']}"
        f"/versions/{other['current_version']['version_uuid']}"
    )
    assert injected.status_code == 404
    assert injected.json()["detail"]["code"] == "DELIVERABLE_VERSION_NOT_FOUND"
    assert outsider.get(
        f"/api/ai/deliverables/{created['deliverable_uuid']}/versions"
    ).status_code == 404


def test_structured_version_diff_tracks_paragraph_table_added_and_removed_blocks(
    client_for_user,
    professional_catalog,
) -> None:
    client = client_for_user("u-1")
    base_content = {
        "schema_version": "1",
        "blocks": [
            {
                "block_id": "overview",
                "type": "paragraph",
                "text": "本月发现两项风险。",
            },
            {
                "block_id": "risk-table",
                "type": "table",
                "rows": [
                    {"row_id": "risk-1", "cells": ["主机风险", "处理中"]},
                ],
            },
            {
                "block_id": "legacy-note",
                "type": "paragraph",
                "text": "旧版临时说明。",
            },
        ],
    }
    updated_content = {
        "schema_version": "1",
        "blocks": [
            {
                "block_id": "overview",
                "type": "paragraph",
                "text": "本月发现两项风险，均已闭环。",
            },
            {
                "block_id": "risk-table",
                "type": "table",
                "rows": [
                    {"row_id": "risk-1", "cells": ["主机风险", "已关闭"]},
                    {"row_id": "risk-2", "cells": ["账号风险", "已关闭"]},
                ],
            },
            {
                "block_id": "monthly-actions",
                "type": "paragraph",
                "text": "完成两项整改。",
            },
        ],
    }
    created = _create(
        client,
        professional_catalog,
        key="create-diff-report",
        content=base_content,
    ).json()
    version_1_uuid = created["current_version"]["version_uuid"]
    version_2_response = _create_version(
        client,
        created["deliverable_uuid"],
        key="create-diff-report-v2",
        content=updated_content,
        change_summary="闭环风险并删除临时说明",
    )
    assert version_2_response.status_code == 201, version_2_response.text
    version_2_uuid = version_2_response.json()["version"]["version_uuid"]

    response = client.get(
        f"/api/ai/deliverables/{created['deliverable_uuid']}/diff",
        params={"from": version_1_uuid, "to": version_2_uuid},
    )

    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["request_id"]
    assert payload["deliverable_uuid"] == created["deliverable_uuid"]
    assert payload["from_version_uuid"] == version_1_uuid
    assert payload["from_version_no"] == 1
    assert payload["to_version_uuid"] == version_2_uuid
    assert payload["to_version_no"] == 2
    assert payload["summary"] == {
        "added_blocks": 1,
        "removed_blocks": 1,
        "modified_blocks": 2,
        "unchanged_blocks": 0,
    }
    assert [change["block_id"] for change in payload["changes"]] == [
        "overview",
        "risk-table",
        "monthly-actions",
        "legacy-note",
    ]
    changes = {change["block_id"]: change for change in payload["changes"]}
    assert changes["overview"]["change_type"] == "modified"
    assert changes["overview"]["field_changes"] == [
        {
            "path": "/text",
            "change_type": "modified",
            "before": "本月发现两项风险。",
            "after": "本月发现两项风险，均已闭环。",
        }
    ]
    assert changes["risk-table"]["change_type"] == "modified"
    assert {
        change["path"] for change in changes["risk-table"]["field_changes"]
    } == {"/rows/0/cells/1", "/rows/1"}
    assert changes["monthly-actions"]["change_type"] == "added"
    assert changes["monthly-actions"]["before"] is None
    assert changes["legacy-note"]["change_type"] == "removed"
    assert changes["legacy-note"]["after"] is None


def test_project_version_reads_allow_read_only_member_and_hide_from_outsider(
    client_for_user,
    professional_catalog,
) -> None:
    owner = client_for_user("u-1")
    read_only = client_for_user("u-2")
    outsider = client_for_user("u-3")
    project = owner.post(
        "/api/ai/projects",
        json={"name": "版本读取隔离项目", "description": ""},
    ).json()
    owner.post(
        f"/api/ai/projects/{project['project_uuid']}/members",
        json={"user_id": "u-2", "role": "read_only"},
    )
    created = _create(
        owner,
        professional_catalog,
        key="create-project-version-read-report",
        scope_type="project",
        formality="formal",
        project_uuid=project["project_uuid"],
    ).json()
    version_uuid = created["current_version"]["version_uuid"]
    history_path = f"/api/ai/deliverables/{created['deliverable_uuid']}/versions"
    version_path = f"{history_path}/{version_uuid}"
    diff_path = f"/api/ai/deliverables/{created['deliverable_uuid']}/diff"

    assert read_only.get(history_path).status_code == 200
    assert read_only.get(version_path).status_code == 200
    assert read_only.get(
        diff_path,
        params={"from": version_uuid, "to": version_uuid},
    ).status_code == 200

    assert outsider.get(history_path).status_code == 404
    assert outsider.get(version_path).status_code == 404
    assert outsider.get(
        diff_path,
        params={"from": version_uuid, "to": version_uuid},
    ).status_code == 404

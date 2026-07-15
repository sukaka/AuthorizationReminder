from datetime import datetime
from io import BytesIO
from types import SimpleNamespace

import pytest
from docx import Document
from sqlalchemy import func, select


CONTENT = {
    "schema_version": "1",
    "blocks": [
        {
            "block_id": "monthly-overview",
            "type": "paragraph",
            "text": "本月未发生重大安全事件。",
        }
    ],
}


@pytest.fixture
def professional_catalog(generation_db):
    from app.professional_delivery.models import (
        SkillDefinition,
        SkillVersion,
        TemplateDefinition,
        TemplateVersion,
    )

    skill = SkillDefinition(
        skill_key="security_ops_monthly_report",
        name="安全运营月报",
        category="security_operations",
        description="形成安全运营月度专业成果",
        scope_policy="both",
        status="published",
        created_by="system",
    )
    template = TemplateDefinition(
        template_key="security_ops_monthly_report",
        name="安全运营月报模板",
        purpose="安全运营月度汇报",
        deliverable_types_json=["security_ops_monthly_report"],
        scope_type="system",
        status="published",
        created_by="system",
    )
    generation_db.add_all([skill, template])
    generation_db.flush()
    template_version = TemplateVersion(
        template_id=template.id,
        version=1,
        content_hash="t" * 64,
        input_schema_json={"type": "object"},
        structure_dsl_json={"type": "document", "children": []},
        word_render_config_json={"template_name": "juxin_standard"},
        status="published",
        published_by="system",
        published_at=datetime(2026, 7, 14, 9, 0, 0),
        created_by="system",
    )
    generation_db.add(template_version)
    generation_db.flush()
    skill_version = SkillVersion(
        skill_id=skill.id,
        version=1,
        content_hash="s" * 64,
        input_schema_json={"type": "object"},
        output_schema_json={"type": "object"},
        plan_definition_json={"steps": []},
        default_template_version_id=template_version.id,
        status="published",
        published_by="system",
        published_at=datetime(2026, 7, 14, 9, 0, 0),
        created_by="system",
    )
    generation_db.add(skill_version)
    generation_db.flush()
    skill.current_published_version_id = skill_version.id
    template.current_published_version_id = template_version.id
    generation_db.commit()
    return SimpleNamespace(
        skill_version=skill_version,
        template_version=template_version,
    )


@pytest.fixture
def export_storage(tmp_path):
    from app.config import get_settings
    from app.main import app

    settings = get_settings().model_copy(
        update={"export_storage_dir": str(tmp_path / "exports")}
    )
    app.dependency_overrides[get_settings] = lambda: settings
    try:
        yield tmp_path / "exports"
    finally:
        app.dependency_overrides.pop(get_settings, None)


def _create(client, catalog, *, key: str) -> dict:
    response = client.post(
        "/api/ai/deliverables",
        headers={"Idempotency-Key": key},
        json={
            "title": "2026 年 7 月安全运营月报",
            "deliverable_type": "security_ops_monthly_report",
            "scope_type": "personal",
            "formality": "formal",
            "skill_version_uuid": catalog.skill_version.uuid,
            "template_version_uuid": catalog.template_version.uuid,
            "content": CONTENT,
            "content_summary": "安全运营月度概览",
            "creation_reason": "manual",
        },
    )
    assert response.status_code == 201, response.text
    return response.json()


def _artifact_and_version(generation_db, created: dict):
    from app.models import WorkArtifact, WorkArtifactVersion

    artifact = generation_db.scalar(
        select(WorkArtifact).where(
            WorkArtifact.uuid == created["deliverable_uuid"]
        )
    )
    assert artifact is not None
    version = generation_db.get(WorkArtifactVersion, artifact.current_version_id)
    assert version is not None
    return artifact, version


def _export_body(artifact, version, **overrides) -> dict:
    body = {
        "row_version": artifact.row_version,
        "content_hash": version.content_hash,
        "export_format": "docx",
    }
    body.update(overrides)
    return body


def _export(client, artifact, version, *, key: str, body: dict | None = None):
    return client.post(
        f"/api/ai/deliverables/{artifact.uuid}/versions/{version.uuid}/exports",
        headers={"Idempotency-Key": key},
        json=body or _export_body(artifact, version),
    )


def _docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join([*paragraphs, *cells])


def test_draft_export_is_watermarked_downloadable_and_idempotent(
    client_for_user,
    generation_db,
    professional_catalog,
    export_storage,
) -> None:
    from app.governance_models import AuditLog
    from app.models import WorkArtifactVersion
    from app.professional_delivery.models import DeliverableExport

    owner = client_for_user("u-1")
    outsider = client_for_user("u-2")
    created = _create(owner, professional_catalog, key="export-draft-create")
    artifact, version = _artifact_and_version(generation_db, created)
    body = _export_body(artifact, version)

    response = _export(
        owner,
        artifact,
        version,
        key="export-draft",
        body=body,
    )
    assert response.status_code == 201, response.text
    payload = response.json()
    assert payload["deliverable_uuid"] == artifact.uuid
    assert payload["version_uuid"] == version.uuid
    assert payload["content_hash"] == version.content_hash
    assert payload["status"] == "ready"
    assert payload["watermarked"] is True
    assert payload["file_name"].endswith("未批准.docx")
    assert len(payload["file_hash"]) == 64
    assert payload["file_size"] > 0
    assert payload["renderer_version"] == "professional-docx-v1"
    assert payload["download_url"] == (
        f"/api/ai/deliverable-exports/{payload['export_uuid']}/download"
    )
    assert generation_db.scalar(
        select(func.count(WorkArtifactVersion.id)).where(
            WorkArtifactVersion.artifact_id == artifact.id
        )
    ) == 1

    downloaded = owner.get(payload["download_url"])
    assert downloaded.status_code == 200, downloaded.text
    assert downloaded.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    text = _docx_text(downloaded.content)
    assert "未批准" in text
    assert "本月未发生重大安全事件" in text
    assert outsider.get(payload["download_url"]).status_code == 404

    replay = _export(
        owner,
        artifact,
        version,
        key="export-draft",
        body=body,
    )
    assert replay.status_code == 201, replay.text
    assert replay.json()["export_uuid"] == payload["export_uuid"]
    assert generation_db.scalar(select(func.count(DeliverableExport.id))) == 1
    assert generation_db.scalar(
        select(func.count(AuditLog.id)).where(
            AuditLog.action == "professional_deliverable.export.create"
        )
    ) == 1
    assert len(list(export_storage.glob("*.docx"))) == 1

    reused = _export(
        owner,
        artifact,
        version,
        key="export-draft",
        body={**body, "content_hash": "f" * 64},
    )
    assert reused.status_code == 409
    assert reused.json()["detail"]["code"] == "IDEMPOTENCY_KEY_REUSED"


def test_approved_exact_export_is_unwatermarked_and_can_be_delivered(
    client_for_user,
    generation_db,
    professional_catalog,
    export_storage,
) -> None:
    from app.models import WorkArtifact, WorkArtifactVersion

    owner = client_for_user("u-1")
    created = _create(owner, professional_catalog, key="export-approved-create")
    artifact, version = _artifact_and_version(generation_db, created)
    artifact.lifecycle_status = "approved"
    artifact.approved_version_id = version.id
    artifact.approved_content_hash = version.content_hash
    artifact.row_version += 1
    generation_db.commit()

    response = _export(owner, artifact, version, key="export-approved")
    assert response.status_code == 201, response.text
    payload = response.json()
    assert payload["watermarked"] is False
    assert "未批准" not in payload["file_name"]
    downloaded = owner.get(payload["download_url"])
    assert downloaded.status_code == 200
    assert "未批准" not in _docx_text(downloaded.content)

    delivered = owner.post(
        f"/api/ai/deliverables/{artifact.uuid}/deliver",
        headers={"Idempotency-Key": "deliver-approved-export"},
        json={
            "row_version": artifact.row_version,
            "version_uuid": version.uuid,
            "content_hash": version.content_hash,
            "export_uuid": payload["export_uuid"],
            "recipient_description": "客户安全负责人",
            "note": "经线下确认后人工交付",
        },
    )
    assert delivered.status_code == 201, delivered.text
    assert delivered.json()["lifecycle_status"] == "delivered"
    assert delivered.json()["delivery"]["export_uuid"] == payload["export_uuid"]
    assert generation_db.scalar(
        select(func.count(WorkArtifactVersion.id)).where(
            WorkArtifactVersion.artifact_id == artifact.id
        )
    ) == 1
    generation_db.refresh(artifact)
    assert artifact.delivered_version_id == version.id
    assert generation_db.get(WorkArtifact, artifact.id) is not None
    assert len(list(export_storage.glob("*.docx"))) == 1


def test_export_rejects_stale_wrong_hash_and_cross_deliverable_version(
    client_for_user,
    generation_db,
    professional_catalog,
    export_storage,
) -> None:
    owner = client_for_user("u-1")
    first = _create(owner, professional_catalog, key="export-target-first")
    second = _create(owner, professional_catalog, key="export-target-second")
    first_artifact, first_version = _artifact_and_version(generation_db, first)
    _, second_version = _artifact_and_version(generation_db, second)

    stale = _export(
        owner,
        first_artifact,
        first_version,
        key="export-target-stale",
        body=_export_body(
            first_artifact,
            first_version,
            row_version=first_artifact.row_version + 1,
        ),
    )
    assert stale.status_code == 409
    assert stale.json()["detail"]["code"] == "DELIVERABLE_VERSION_CONFLICT"

    wrong_hash = _export(
        owner,
        first_artifact,
        first_version,
        key="export-target-hash",
        body=_export_body(
            first_artifact,
            first_version,
            content_hash="f" * 64,
        ),
    )
    assert wrong_hash.status_code == 409
    assert wrong_hash.json()["detail"]["code"] == "DELIVERABLE_TARGET_MISMATCH"

    cross_version = _export(
        owner,
        first_artifact,
        second_version,
        key="export-target-cross",
        body=_export_body(first_artifact, second_version),
    )
    assert cross_version.status_code == 404
    assert cross_version.json()["detail"]["code"] == "DELIVERABLE_VERSION_NOT_FOUND"
    assert list(export_storage.glob("*.docx")) == []


def test_export_audit_failure_rolls_back_record_and_removes_file(
    client_for_user,
    generation_db,
    professional_catalog,
    export_storage,
    monkeypatch,
) -> None:
    from app.professional_delivery import routes
    from app.professional_delivery.models import DeliverableExport

    owner = client_for_user("u-1")
    created = _create(owner, professional_catalog, key="export-rollback-create")
    artifact, version = _artifact_and_version(generation_db, created)

    def fail_audit(*_args, **_kwargs) -> None:
        raise RuntimeError("audit unavailable")

    monkeypatch.setattr(routes, "write_request_audit", fail_audit)
    with pytest.raises(RuntimeError, match="audit unavailable"):
        _export(owner, artifact, version, key="export-rollback")

    generation_db.expire_all()
    assert generation_db.scalar(select(func.count(DeliverableExport.id))) == 0
    assert list(export_storage.glob("*.docx")) == []

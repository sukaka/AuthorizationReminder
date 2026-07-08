from io import BytesIO

from docx import Document
from docx.oxml.ns import qn

from app.document_templates.base import DocumentRenderPayload
from app.document_templates.registry import get_document_template
from app.document_templates.structure_validator import strip_duplicate_template_headings
from app.word_export import COMPANY_WORD_STYLE, render_generation_docx


def test_document_template_registry_returns_fallback_for_unknown_code():
    template = get_document_template("")

    assert template.code == "generic_v1"
    assert template.name == "通用正式文档模板"


def test_structure_validator_strips_duplicate_company_headings():
    cleaned = strip_duplicate_template_headings(
        "# 一、任务说明\n\n正文\n\n# 二、基本信息\n\n重复内容\n\n# 三、背景说明\n\n重复背景",
        fixed_headings=("基本信息", "背景说明"),
    )

    assert "任务说明" in cleaned
    assert "# 二、基本信息" not in cleaned
    assert "# 三、背景说明" not in cleaned
    assert "重复内容" in cleaned
    assert "重复背景" in cleaned


def test_document_template_render_keeps_single_fixed_heading_with_user_content():
    template = get_document_template("generic_v1")

    payload = template.render_docx(
        DocumentRenderPayload(
            title="客户项目交付报告",
            task_name="项目交付",
            department="产品交付部",
            author="张三",
            output="# 基本信息\n\n客户名称：A",
            version="1.0.0",
        )
    )

    document = Document(BytesIO(payload))
    headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style.name in {"Heading 1", "Heading 2", "Heading 3"}
    ]
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert sum(heading.endswith("基本信息") for heading in headings) == 1
    assert "客户名称：A" in text


def test_work_plan_template_renders_without_duplicate_fixed_headings():
    template = get_document_template("work_plan_v1")
    payload = DocumentRenderPayload(
        title="阶段工作计划",
        task_name="工作计划",
        department="产品交付部",
        author="张三",
        version="1.0.0",
        output="# 一、任务说明\n\n正文\n\n# 二、基本信息\n\n重复基本信息",
    )

    document = Document(BytesIO(template.render_docx(payload)))
    headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style.name in {"Heading 1", "Heading 2", "Heading 3"}
    ]
    basic_info_headings = [heading for heading in headings if heading.endswith("基本信息")]

    assert template.name == "阶段工作计划模板"
    assert len(basic_info_headings) == 1
    assert any(heading.endswith("工作目标与范围") for heading in headings)


def test_work_plan_template_preserves_content_under_canonical_heading():
    template = get_document_template("work_plan_v1")

    normalized = template.normalize_output("# 二、基本信息\n\n客户名称：A")

    assert "# 基本信息\n\n客户名称：A" in normalized
    assert normalized.index("# 基本信息") < normalized.index("客户名称：A")
    assert normalized.index("客户名称：A") < normalized.index("# 背景说明")


def test_work_plan_template_keeps_nested_heading_inside_fixed_section():
    template = get_document_template("work_plan_v1")

    normalized = template.normalize_output("# 基本信息\n\n## 客户信息\n\n客户名称：A")

    assert normalized.index("# 基本信息") < normalized.index("## 客户信息")
    assert normalized.index("## 客户信息") < normalized.index("客户名称：A")
    assert normalized.index("客户名称：A") < normalized.index("# 背景说明")


def test_meeting_minutes_template_renders_action_item_table():
    template = get_document_template("meeting_minutes_v1")
    payload = DocumentRenderPayload(
        title="会议纪要",
        task_name="会议纪要",
        department="产品交付部",
        author="张三",
        version="1.0.0",
        output="会议讨论了项目进度。",
    )

    document = Document(BytesIO(template.render_docx(payload)))
    table_texts = [
        [cell.text for cell in row.cells]
        for table in document.tables
        for row in table.rows
    ]

    assert ["序号", "事项", "责任人", "截止时间", "状态", "备注"] in table_texts


def test_meeting_minutes_template_places_default_table_in_action_item_section():
    template = get_document_template("meeting_minutes_v1")

    normalized = template.normalize_output("会议讨论了项目进度。")

    assert normalized.index("# 待办事项表") < normalized.index("| 序号 | 事项 | 责任人 | 截止时间 | 状态 | 备注 |")
    assert normalized.index("| 序号 | 事项 | 责任人 | 截止时间 | 状态 | 备注 |") < normalized.index("# 风险与分歧")


def test_meeting_minutes_template_preserves_action_item_content_before_default_table():
    template = get_document_template("meeting_minutes_v1")

    normalized = template.normalize_output("# 待办事项表\n\n请张三跟进客户")

    assert normalized.index("# 待办事项表") < normalized.index("请张三跟进客户")
    assert normalized.index("请张三跟进客户") < normalized.index("| 序号 | 事项 | 责任人 | 截止时间 | 状态 | 备注 |")
    assert normalized.index("| 序号 | 事项 | 责任人 | 截止时间 | 状态 | 备注 |") < normalized.index("# 风险与分歧")


def test_project_report_template_is_registered():
    template = get_document_template("project_report_v1")

    assert template.name == "项目汇报模板"


def test_company_word_style_constants_are_named():
    assert COMPANY_WORD_STYLE["page"]["top_margin_cm"] == 2.5
    assert COMPANY_WORD_STYLE["page"]["left_margin_cm"] == 2.8
    assert COMPANY_WORD_STYLE["brand"]["header_line_color"] in {"C00000", "D9D9D9"}
    assert "基本信息" in COMPANY_WORD_STYLE["required_sections"]


def test_render_word_uses_v110_page_and_brand_rules():
    payload = render_generation_docx(
        title="项目实施报告",
        task_name="实施报告",
        department="产品交付部",
        author="张三",
        output="# 一、项目背景\n\n正文\n\n| 项目 | 内容 |\n|---|---|\n| 状态 | 已完成 |",
        version="V1.0",
    )

    document = Document(BytesIO(payload))
    section = document.sections[0]

    assert round(section.top_margin.cm, 1) == 2.5
    assert round(section.left_margin.cm, 1) == 2.8
    assert "聚信得仁" in section.header.paragraphs[0].text
    assert any(table.cell(0, 0).text == "项目" for table in document.tables)


def test_render_word_numbers_headings_by_level():
    payload = render_generation_docx(
        title="项目实施报告",
        task_name="实施报告",
        department="产品交付部",
        author="张三",
        output=(
            "# 项目背景\n\n"
            "## 建设范围\n\n"
            "### 系统边界\n\n"
            "## 交付内容\n\n"
            "# 实施计划\n\n"
            "### 注意事项\n\n"
            "# 三、已有编号"
        ),
        version="V1.0",
    )

    document = Document(BytesIO(payload))
    headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style.name in {"Heading 1", "Heading 2", "Heading 3"}
    ]

    assert headings[:7] == [
        "一、修订记录",
        "二、项目背景",
        "1. 建设范围",
        "（1）系统边界",
        "2. 交付内容",
        "三、实施计划",
        "（1）注意事项",
    ]
    assert "四、已有编号" in headings
    assert "四、三、已有编号" not in headings


def test_render_word_renumbers_existing_heading_prefixes():
    payload = render_generation_docx(
        title="项目实施报告",
        task_name="实施报告",
        department="产品交付部",
        author="张三",
        output=(
            "# 一、项目背景\n\n"
            "# 实施计划\n\n"
            "## 2. 建设范围\n\n"
            "## 交付内容\n\n"
            "### （3）系统边界\n\n"
            "### 验收要求"
        ),
        version="V1.0",
    )

    document = Document(BytesIO(payload))
    headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style.name in {"Heading 1", "Heading 2", "Heading 3"}
    ]

    assert headings[:7] == [
        "一、修订记录",
        "二、项目背景",
        "三、实施计划",
        "1. 建设范围",
        "2. 交付内容",
        "（1）系统边界",
        "（2）验收要求",
    ]
    assert "一、项目背景" not in headings
    assert "2. 2. 建设范围" not in headings


def test_render_word_cleans_markdown_markers_from_business_output():
    payload = render_generation_docx(
        title="客户评分建议",
        task_name="客户评分",
        department="销售部",
        author="张三",
        output=(
            "**75分（满分100分）**\n\n"
            "**8. 评分依据**\n\n"
            "* **行业匹配度（20/20）**：网络安全行业，客户群匹配。\n"
            "* **预算可行性（10/20）**：预算紧张。\n\n"
            "---\n\n"
            "## 9. 建议跟进级别\n\n"
            "**A级（重点跟进）**"
        ),
        version="V1.0",
    )

    document = Document(BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "75分（满分100分）" in text
    assert "8. 评分依据" in text
    assert "行业匹配度（20/20）：网络安全行业，客户群匹配。" in text
    assert "预算可行性（10/20）：预算紧张。" in text
    assert "建议跟进级别" in text
    assert "A级（重点跟进）" in text
    assert "**" not in text
    assert "---" not in text


def test_render_word_preserves_text_while_cleaning_unmatched_markers():
    payload = render_generation_docx(
        title="异常语法报告",
        task_name="异常语法",
        department="产品交付部",
        author="张三",
        output="未闭合 **文本",
        version="V1.0",
    )

    document = Document(BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "未闭合 文本" in text
    assert "**" not in text


def test_render_word_applies_company_output_control_footer_and_brand_line():
    payload = render_generation_docx(
        title="客户项目交付报告",
        task_name="项目交付",
        department="产品交付部",
        author="张三",
        output="# 基本信息\n\n正文",
        version="1.0.0",
    )

    document = Document(BytesIO(payload))
    section = document.sections[0]
    footer_xml = section.footer.paragraphs[0]._p.xml
    header_border = section.header.paragraphs[0]._p.get_or_add_pPr().find(qn("w:pBdr"))
    header_bottom = None if header_border is None else header_border.find(qn("w:bottom"))

    assert "北京聚信得仁科技有限公司" in section.footer.paragraphs[0].text
    assert "客户项目交付文档" in section.footer.paragraphs[0].text
    assert "共 " in section.footer.paragraphs[0].text
    assert "NUMPAGES" in footer_xml
    assert header_bottom is not None
    assert header_bottom.get(qn("w:color")) in {"C00000", "D9D9D9"}


def test_render_word_backfills_company_required_structure_and_final_checks():
    payload = render_generation_docx(
        title="客户项目交付报告",
        task_name="项目交付",
        department="产品交付部",
        author="张三",
        output="# 主要内容\n\n已生成正文",
        version="1.0.0",
    )

    document = Document(BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    for section in (
        "基本信息",
        "背景说明",
        "目标与范围",
        "主要内容",
        "执行步骤或工作安排",
        "表格清单或结果统计",
        "风险与注意事项",
        "需确认事项",
        "交付物或附件",
        "结论与下一步计划",
    ):
        assert section in text

    for item in (
        "待确认事项",
        "需人工复核事项",
        "不建议直接对外发送的内容",
        "可以直接落地执行的下一步动作",
    ):
        assert item in text

    for label in ("已知事实", "合理判断", "风险提醒"):
        assert label in text


def test_render_word_fills_blank_table_cells_with_pending_confirmation():
    payload = render_generation_docx(
        title="客户项目交付报告",
        task_name="项目交付",
        department="产品交付部",
        author="张三",
        output="| 字段 | 内容 |\n|---|---|\n| 客户名称 | |",
        version="1.0.0",
    )

    document = Document(BytesIO(payload))

    data_table = next(
        table
        for table in document.tables
        if any(cell.text == "客户名称" for row in table.rows for cell in row.cells)
    )

    assert data_table.cell(1, 1).text == "待确认"


def test_render_word_supports_markdown_quote_blocks():
    payload = render_generation_docx(
        title="引用块测试",
        task_name="聊天导出",
        department="产品交付部",
        author="张三",
        output="> 客户原话：需要下周完成验收\n> 请保留引用内容",
        version="1.0.0",
    )

    document = Document(BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "客户原话：需要下周完成验收" in text
    assert "请保留引用内容" in text
    assert "> 客户原话" not in text

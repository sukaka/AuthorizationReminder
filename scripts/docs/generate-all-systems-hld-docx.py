#!/usr/bin/env python3
"""Generate the all-systems HLD Word document and embedded architecture diagrams."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs/superpowers/specs/2026-06-13-all-systems-high-level-design.md"
OUTPUT_DIR = ROOT / "outputs/all-systems-hld"
DOCX_PATH = OUTPUT_DIR / "聚信多系统业务平台高层设计.docx"
DOCUMENT_VERSION = "5.70.15"
FONT_CANDIDATES = [
    Path("/System/Library/Fonts/STHeiti Medium.ttc"),
    Path("/System/Library/Fonts/STHeiti Light.ttc"),
    Path("/System/Library/Fonts/PingFang.ttc"),
]
FONT_PATH = next((path for path in FONT_CANDIDATES if path.exists()), None)

NAVY = "#12365A"
BLUE = "#1F6FB2"
CYAN = "#30A6C7"
TEAL = "#168B80"
GREEN = "#2C8D62"
ORANGE = "#D8802A"
PURPLE = "#7157A5"
RED = "#B64747"
INK = "#263646"
MUTED = "#667788"
LIGHT = "#EEF4F8"
PALE_BLUE = "#EAF3FA"
WHITE = "#FFFFFF"


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    if FONT_PATH:
        return ImageFont.truetype(str(FONT_PATH), size=size, index=0)
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> tuple[int, int]:
    box = draw.multiline_textbbox((0, 0), text, font=font, spacing=5, align="center")
    return box[2] - box[0], box[3] - box[1]


def round_rect(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    fill: str,
    outline: str = "#C5D3DE",
    radius: int = 22,
    width: int = 3,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    size: int = 30,
    fill: str = INK,
    bold: bool = False,
) -> None:
    font = pil_font(size, bold=bold)
    width, height = text_size(draw, text, font)
    x1, y1, x2, y2 = box
    draw.multiline_text(
        ((x1 + x2 - width) / 2, (y1 + y2 - height) / 2),
        text,
        font=font,
        fill=fill,
        spacing=5,
        align="center",
    )


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: str = MUTED,
    width: int = 5,
) -> None:
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    angle = 13
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - direction * 18, y2 - angle), (x2 - direction * 18, y2 + angle)]
    else:
        direction = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - angle, y2 - direction * 18), (x2 + angle, y2 - direction * 18)]
    draw.polygon(points, fill=color)


def canvas(title: str, subtitle: str, width: int = 2200, height: int = 1320) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (width, height), WHITE)
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width, 140), fill=NAVY)
    draw.text((70, 32), title, font=pil_font(48, True), fill=WHITE)
    draw.text((72, 92), subtitle, font=pil_font(25), fill="#D9E8F4")
    return image, draw


def save_diagram(image: Image.Image, name: str) -> Path:
    path = OUTPUT_DIR / name
    image.save(path, format="PNG", optimize=True)
    return path


def draw_overview() -> Path:
    image, draw = canvas("聚信多系统业务平台总体架构", "当前真实架构：统一门户 + 11 个业务系统 + 共享基础设施")
    round_rect(draw, (80, 180, 2120, 300), PALE_BLUE, BLUE)
    centered_text(draw, (80, 180, 2120, 300), "访问与平台层：统一认证门户  |  管理中心  |  审计中心", 34, NAVY, True)

    system_names = [
        "授权提醒", "交付系统", "CMDB", "库存管理", "设备流转", "文档管理",
        "标书协同", "培训考试", "提示词中心", "SCA", "统一大屏",
    ]
    colors = [BLUE, TEAL, PURPLE, ORANGE, CYAN, GREEN, PURPLE, BLUE, TEAL, RED, NAVY]
    box_w, box_h, gap = 300, 105, 35
    first_x = 90
    for idx, name in enumerate(system_names):
        row = 0 if idx < 6 else 1
        col = idx if row == 0 else idx - 6
        x = first_x + col * (box_w + gap)
        if row == 1:
            x += 165
        y = 365 + row * 145
        round_rect(draw, (x, y, x + box_w, y + box_h), WHITE, colors[idx], width=4)
        centered_text(draw, (x, y, x + box_w, y + box_h), name, 28, colors[idx], True)

    arrow(draw, (1100, 300), (1100, 350), BLUE)
    round_rect(draw, (80, 690, 2120, 850), LIGHT, "#9BB3C5")
    centered_text(
        draw,
        (80, 690, 2120, 850),
        "业务服务与集成层\nNode.js/Express  |  Go/Gin  |  Python/FastAPI  |  物流网关  |  OnlyOffice  |  扫描引擎",
        30,
        INK,
        True,
    )
    arrow(draw, (1100, 655), (1100, 690), MUTED)

    data_boxes = [
        (100, 960, 570, 1165, "MySQL 8\n多业务 Schema", BLUE),
        (640, 960, 1050, 1165, "PostgreSQL 16\nSCA 主数据", PURPLE),
        (1120, 960, 1530, 1165, "Redis 7\n缓存与 Celery", RED),
        (1600, 960, 2100, 1165, "持久化文件卷\n文档 / 报告 / SBOM", TEAL),
    ]
    for x1, y1, x2, y2, label, color in data_boxes:
        round_rect(draw, (x1, y1, x2, y2), WHITE, color, width=4)
        centered_text(draw, (x1, y1, x2, y2), label, 29, color, True)
    arrow(draw, (1100, 850), (1100, 935), MUTED)
    draw.text((85, 1235), "历史兼容：Ticketing / Sec-Impl 的门户权限键统一归一化为 Delivery", font=pil_font(25), fill=MUTED)
    return save_diagram(image, "01-platform-overview.png")


def draw_auth_flow() -> Path:
    image, draw = canvas("统一认证与系统访问流程", "HttpOnly Cookie、门户授权、业务 API 二次鉴权")
    nodes = [
        ("用户浏览器", 80, BLUE),
        ("Auth 登录\n5180", 420, NAVY),
        ("门户应用清单\n/api/auth/apps", 790, TEAL),
        ("业务前端", 1190, CYAN),
        ("业务 API", 1510, ORANGE),
        ("Auth Introspection\n系统键校验", 1820, PURPLE),
    ]
    y1, y2 = 430, 620
    for label, x, color in nodes:
        round_rect(draw, (x, y1, x + 270, y2), WHITE, color, width=4)
        centered_text(draw, (x, y1, x + 270, y2), label, 27, color, True)
    for index in range(len(nodes) - 1):
        x = nodes[index][1] + 270
        next_x = nodes[index + 1][1]
        arrow(draw, (x + 8, 525), (next_x - 8, 525), MUTED)

    steps = [
        "1. 登录并设置会话 Cookie",
        "2. 获取用户可访问系统",
        "3. 跳转目标业务前端",
        "4. 请求携带统一 Cookie",
        "5. API 校验身份、角色与系统键",
    ]
    start_x = 140
    for idx, step in enumerate(steps):
        y = 760 + idx * 78
        draw.ellipse((start_x, y, start_x + 46, y + 46), fill=BLUE)
        centered_text(draw, (start_x, y, start_x + 46, y + 46), str(idx + 1), 23, WHITE, True)
        draw.text((start_x + 70, y + 6), step, font=pil_font(28), fill=INK)
    round_rect(draw, (1240, 770, 2080, 1140), PALE_BLUE, BLUE)
    centered_text(
        draw,
        (1240, 770, 2080, 1140),
        "授权原则\n\n前端菜单只负责展示\n业务后端必须再次鉴权\n系统键与资源范围同时校验\nAuth 不可用时拒绝受保护操作",
        30,
        NAVY,
        True,
    )
    return save_diagram(image, "02-auth-access-flow.png")


def draw_data_topology() -> Path:
    image, draw = canvas("平台数据拓扑", "单 MySQL 实例多 Schema；SCA 独立 PostgreSQL / Redis；文件卷单独备份")
    round_rect(draw, (80, 200, 1380, 1160), PALE_BLUE, BLUE, width=4)
    centered_text(draw, (110, 220, 520, 300), "MySQL 8 单实例", 38, NAVY, True)
    schemas = [
        ("juxin_reminder", "Auth / Reminder\n历史 Ticketing", BLUE),
        ("juxin_delivery", "Delivery", TEAL),
        ("cmdb", "CMDB", PURPLE),
        ("juxin_inventory", "Inventory", ORANGE),
        ("juxin_device_flow", "Device Flow", CYAN),
        ("juxin_faq", "FAQ\nTrain Exam 附加读取", GREEN),
        ("juxin_tender", "Tender", PURPLE),
        ("juxin_train_exam", "Train Exam", BLUE),
        ("juxin_prompt_center", "Prompt Center", TEAL),
        ("juxin_big_screen", "Big Screen", NAVY),
    ]
    for idx, (schema, owner, color) in enumerate(schemas):
        row, col = divmod(idx, 2)
        x = 120 + col * 610
        y = 330 + row * 155
        round_rect(draw, (x, y, x + 550, y + 115), WHITE, color, width=3)
        draw.text((x + 25, y + 20), schema, font=pil_font(27, True), fill=color)
        draw.text((x + 25, y + 65), owner, font=pil_font(22), fill=INK)

    side_boxes = [
        (1490, 220, 2100, 470, "PostgreSQL 16\njuxin_sca\nSCA + Dependency-Track", PURPLE),
        (1490, 535, 2100, 755, "Redis 7\n缓存 / Broker / 结果", RED),
        (1490, 820, 2100, 1120, "文件与对象数据\nFAQ / Tender / Train Exam\nSCA 报告 / SBOM / 扫描结果", TEAL),
    ]
    for x1, y1, x2, y2, text, color in side_boxes:
        round_rect(draw, (x1, y1, x2, y2), WHITE, color, width=4)
        centered_text(draw, (x1, y1, x2, y2), text, 29, color, True)
    arrow(draw, (1380, 630), (1470, 630), MUTED)
    draw.text((100, 1230), "治理重点：独立最小权限账号、跨库直连改 API、数据库与文件卷配套恢复", font=pil_font(26), fill=MUTED)
    return save_diagram(image, "03-data-topology.png")


def draw_deployment() -> Path:
    image, draw = canvas("当前 Docker Compose 部署关系", "单 Compose 网络，前端静态站点、业务 API 与共享中间件分层运行")
    round_rect(draw, (70, 185, 2130, 1190), "#F7FAFC", "#9BB3C5", width=4)
    draw.text((100, 205), "Docker Compose 网络", font=pil_font(32, True), fill=NAVY)

    round_rect(draw, (130, 300, 450, 460), WHITE, BLUE, width=4)
    centered_text(draw, (130, 300, 450, 460), "用户浏览器\n门户 5180", 30, BLUE, True)
    arrow(draw, (460, 380), (580, 380), BLUE)

    round_rect(draw, (590, 260, 1080, 500), PALE_BLUE, NAVY, width=4)
    centered_text(draw, (590, 260, 1080, 500), "接入层\nAuth Portal\n11 套 Web / Nginx\n宿主机端口 18080-18092 / 8090", 29, NAVY, True)
    arrow(draw, (1090, 380), (1210, 380), MUTED)

    round_rect(draw, (1220, 240, 2050, 520), WHITE, TEAL, width=4)
    centered_text(
        draw,
        (1220, 240, 2050, 520),
        "服务层\nReminder / Delivery / CMDB / Inventory\nDevice / FAQ / Tender / Train Exam\nPrompt / SCA / Big Screen",
        28,
        TEAL,
        True,
    )

    infra = [
        (160, 690, 560, 900, "MySQL 8\n53308 -> 3306", BLUE),
        (630, 690, 1030, 900, "PostgreSQL 16\n55433 -> 5432", PURPLE),
        (1100, 690, 1500, 900, "Redis 7\n56380 -> 6379", RED),
        (1570, 690, 2020, 900, "OnlyOffice × 2\n内部服务", ORANGE),
    ]
    for x1, y1, x2, y2, text, color in infra:
        round_rect(draw, (x1, y1, x2, y2), WHITE, color, width=4)
        centered_text(draw, (x1, y1, x2, y2), text, 28, color, True)
    arrow(draw, (1600, 520), (1600, 660), MUTED)

    round_rect(draw, (320, 1000, 1880, 1135), LIGHT, "#9BB3C5")
    centered_text(
        draw,
        (320, 1000, 1880, 1135),
        "异步与集成：SCA Worker / Scanner Worker / Beat  |  Dependency-Track  |  Shipping Gateway  |  持久化卷",
        27,
        INK,
        True,
    )
    draw.text((100, 1235), "生产建议：统一 HTTPS 入口、数据库端口不暴露公网、密钥从环境或密钥系统注入", font=pil_font(26), fill=MUTED)
    return save_diagram(image, "04-compose-deployment.png")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill.replace("#", ""))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name: str = "STHeiti", size: float | None = None, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def add_inline(paragraph, text: str, size: float = 10.5, color: str = INK) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            set_run_font(run, size=size, color=color)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Courier New", size=size - 0.5, color=NAVY)
            run.font.highlight_color = None
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, color=color)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "打开 Word 后右键目录并选择“更新域”。"
    separate.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, separate, end])


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "STHeiti"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25

    heading_colors = {1: NAVY, 2: BLUE, 3: TEAL}
    heading_sizes = {1: 20, 2: 15, 3: 12}
    for level in (1, 2, 3):
        style = styles[f"Heading {level}"]
        style.font.name = "STHeiti"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
        style.font.size = Pt(heading_sizes[level])
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(heading_colors[level].replace("#", ""))
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in [style.name for style in styles]:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Courier New"
        code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        code_style.font.size = Pt(8.5)
        code_style.paragraph_format.left_indent = Cm(0.5)
        code_style.paragraph_format.right_indent = Cm(0.5)
        code_style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"聚信多系统业务平台高层设计  |  v{DOCUMENT_VERSION}")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_cover(document: Document, toc_entries: list[str]) -> None:
    for _ in range(5):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("聚信多系统业务平台")
    set_run_font(run, size=32, color=NAVY)
    run.bold = True
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("高层设计文档")
    set_run_font(run, size=25, color=BLUE)
    run.bold = True
    line = document.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    set_run_font(run, size=13, color=CYAN)
    document.add_paragraph()
    metadata = [
        "文档状态：已确认设计",
        "编制日期：2026-06-13",
        f"适用版本：{DOCUMENT_VERSION} 及后续兼容版本",
        "主要读者：技术负责人、研发、测试、部署与运维交接人员",
    ]
    for item in metadata:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(item)
        set_run_font(run, size=11.5, color=INK)
    for _ in range(5):
        document.add_paragraph()
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("现状基准：根 docker-compose.yml、auth/portal-routing.js 与各系统当前实现")
    set_run_font(run, size=9.5, color=MUTED)
    document.add_page_break()

    heading = document.add_paragraph()
    run = heading.add_run("目录")
    set_run_font(run, size=22, color=NAVY)
    run.bold = True
    table = document.add_table(rows=(len(toc_entries) + 1) // 2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    split = (len(toc_entries) + 1) // 2
    columns = [toc_entries[:split], toc_entries[split:]]
    for col_index, entries in enumerate(columns):
        for row_index, entry in enumerate(entries):
            cell = table.cell(row_index, col_index)
            cell.width = Cm(8)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(5)
            run = paragraph.add_run(entry)
            set_run_font(run, size=10.5, color=INK)
    document.add_page_break()


def add_diagram(document: Document, path: Path, caption: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(6.55))
    caption_para = document.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_para.add_run(caption)
    set_run_font(run, size=9, color=MUTED)
    run.italic = True


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, values in enumerate(rows):
        row = table.add_row()
        if row_index == 0:
            set_repeat_table_header(row)
        for col_index in range(columns):
            cell = row.cells[col_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            value = values[col_index] if col_index < len(values) else ""
            paragraph = cell.paragraphs[0]
            add_inline(paragraph, value, size=8.5 if columns >= 6 else 9)
            if row_index == 0:
                set_cell_shading(cell, NAVY)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            elif row_index % 2 == 0:
                set_cell_shading(cell, "F4F7FA")
    document.add_paragraph()


def add_markdown(document: Document, markdown: str, diagrams: dict[str, Path]) -> None:
    lines = markdown.splitlines()
    index = 0
    mermaid_index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if index == 0 and stripped.startswith("# "):
            index += 1
            continue
        if stripped.startswith("> "):
            index += 1
            continue
        if not stripped:
            index += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip()
            block = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                block.append(lines[index])
                index += 1
            index += 1
            if language == "mermaid":
                key = "overview" if mermaid_index == 0 else "auth"
                caption = "图 1  平台总体架构" if mermaid_index == 0 else "图 4  统一认证与系统访问流程"
                add_diagram(document, diagrams[key], caption)
                mermaid_index += 1
            else:
                paragraph = document.add_paragraph(style="Code Block")
                set_cell_shading_like_paragraph(paragraph, "F2F4F7")
                run = paragraph.add_run("\n".join(block))
                set_run_font(run, name="Courier New", size=8.5, color=INK)
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and lines[index + 1].strip().startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(document, parse_table(table_lines))
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            if level == 2:
                document.add_page_break()
            paragraph = document.add_paragraph(text, style=f"Heading {min(level, 3)}")
            if text.startswith("7. 数据架构"):
                add_diagram(document, diagrams["data"], "图 2  平台数据拓扑")
            elif text.startswith("9. 部署架构"):
                add_diagram(document, diagrams["deployment"], "图 3  当前 Docker Compose 部署关系")
            index += 1
            continue

        bullet_match = re.match(r"^(\s*)-\s+(.+)$", line)
        if bullet_match:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Cm(0.7 + len(bullet_match.group(1)) * 0.2)
            add_inline(paragraph, bullet_match.group(2))
            index += 1
            continue

        numbered_match = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if numbered_match:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.left_indent = Cm(0.7 + len(numbered_match.group(1)) * 0.2)
            add_inline(paragraph, numbered_match.group(2))
            index += 1
            continue

        paragraph = document.add_paragraph()
        add_inline(paragraph, stripped)
        index += 1


def set_cell_shading_like_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def build_document(diagrams: dict[str, Path]) -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    toc_entries = re.findall(r"^##\s+(.+)$", markdown, flags=re.MULTILINE)
    document = Document()
    configure_document(document)
    add_cover(document, toc_entries)
    add_markdown(document, markdown, diagrams)
    properties = document.core_properties
    properties.title = "聚信多系统业务平台高层设计"
    properties.subject = "统一认证门户、11 个业务系统、共享基础设施与架构演进"
    properties.author = "聚信项目组"
    properties.keywords = "高层设计, 系统架构, SSO, Docker Compose, 运维交接"
    document.save(DOCX_PATH)


def main() -> int:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if not SOURCE.exists():
        print(f"Source Markdown not found: {SOURCE}", file=sys.stderr)
        return 1
    diagrams = {
        "overview": draw_overview(),
        "auth": draw_auth_flow(),
        "data": draw_data_topology(),
        "deployment": draw_deployment(),
    }
    build_document(diagrams)
    print(DOCX_PATH)
    for path in diagrams.values():
        print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
